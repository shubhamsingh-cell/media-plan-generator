#!/usr/bin/env python3
"""Generate v3.1 System State DOCX document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # navy

# Section margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.style = doc.styles["Normal"]
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacer
    return table


def add_body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Media Plan Generator v3.1")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("System State Document")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0x8A)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Generated: 2026-03-09\n").font.size = Pt(11)
meta.add_run("Deployment: https://media-plan-generator.onrender.com/\n").font.size = Pt(11)
meta.add_run("Stack: Python stdlib HTTP server | openpyxl + python-pptx | Render.com Standard").font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. DATA SOURCE MATRIX
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Data Source Matrix", level=1)

# A. KB Files
doc.add_heading("A. Knowledge Base Files (13 Static KB)", level=2)
add_body("Total KB size: ~1.0 MB of curated recruitment intelligence. All files validated as parseable JSON.")

kb_data = [
    ("1", "recruitment_industry_knowledge.json", "100.9 KB", "Industry-specific hiring patterns, funnel rates"),
    ("2", "platform_intelligence_deep.json", "84.2 KB", "Deep platform data (CPC/CPA/CTR by platform)"),
    ("3", "recruitment_benchmarks_deep.json", "63.9 KB", "Industry benchmarks, conversion rates"),
    ("4", "channels_db.json", "58.5 KB", "Channel taxonomy, capabilities, pricing models"),
    ("5", "joveo_2026_benchmarks.json", "11.2 KB", "Joveo internal CPA/CPC benchmarks by occupation"),
    ("6", "global_supply.json", "107.2 KB", "Talent supply by country/region"),
    ("7", "industry_white_papers.json", "75.3 KB", "Industry intelligence, strategic insights"),
    ("8", "joveo_publishers.json", "96.4 KB", "Joveo publisher network (Indeed, ZR, etc.)"),
    ("9", "linkedin_guidewire_data.json", "27.9 KB", "LinkedIn audience size and targeting data"),
    ("10", "recruitment_strategy_intelligence.json", "72.0 KB", "Strategy playbooks by collar/tier"),
    ("11", "regional_hiring_intelligence.json", "144.4 KB", "Regional labor market intelligence"),
    ("12", "supply_ecosystem_intelligence.json", "89.7 KB", "Talent pipeline ecosystem data"),
    ("13", "workforce_trends_intelligence.json", "72.4 KB", "Macro workforce trend data"),
]
add_table(["#", "KB File", "Size", "Content"], kb_data, [0.3, 2.8, 0.7, 3.0])

# B. Live APIs
doc.add_heading("B. Live API Sources (33 APIs)", level=2)

doc.add_heading("Labor Market & Economics (13 APIs)", level=3)
api_labor = [
    ("1", "fetch_salary_data", "BLS OES", "Salary ranges by role/location"),
    ("2", "fetch_industry_employment", "BLS CES", "Industry employment levels"),
    ("3", "fetch_bls_jolts", "BLS JOLTS", "Job openings, hires, quits rates"),
    ("4", "fetch_fred_indicators", "FRED (Federal Reserve)", "Unemployment, CPI, wage growth"),
    ("5", "fetch_fred_employment_series", "FRED", "Employment time series"),
    ("6", "fetch_onet_occupation_data", "O*NET", "Skills, education, outlook by SOC"),
    ("7", "fetch_onet_job_zone", "O*NET", "Job zone/complexity level"),
    ("8", "fetch_h1b_wage_benchmarks", "H-1B Disclosure", "Prevailing wage data"),
    ("9", "fetch_careeronestop_data", "CareerOneStop (DOL)", "Certifications, licensing"),
    ("10", "fetch_imf_indicators", "IMF", "Country economic indicators"),
    ("11", "fetch_eurostat_labour_data", "Eurostat", "EU labor statistics"),
    ("12", "fetch_ilo_labour_data", "ILO (UN)", "International labor indicators"),
    ("13", "fetch_global_indicators", "World Bank", "GDP, unemployment, labor force"),
]
add_table(["#", "Function", "Source", "Data"], api_labor, [0.3, 2.3, 1.8, 2.4])

doc.add_heading("Location & Demographics (6 APIs)", level=3)
api_location = [
    ("14", "fetch_location_demographics", "US Census", "Population, education, income"),
    ("15", "fetch_country_data", "REST Countries", "Country demographics"),
    ("16", "fetch_geonames_data", "GeoNames", "City population, coordinates"),
    ("17", "fetch_teleport_city_data", "Teleport", "Quality of life scores"),
    ("18", "fetch_datausa_occupation_stats", "DataUSA", "Occupation employment stats"),
    ("19", "fetch_datausa_location_data", "DataUSA", "Location employment data"),
]
add_table(["#", "Function", "Source", "Data"], api_location, [0.3, 2.3, 1.8, 2.4])

doc.add_heading("Advertising Platforms (5 APIs)", level=3)
api_ads = [
    ("20", "fetch_google_ads_data", "Google Ads API", "CPC, impressions, competition"),
    ("21", "fetch_meta_ads_data", "Meta Marketing API", "CPM, audience size, targeting"),
    ("22", "fetch_bing_ads_data", "Bing Ads API", "CPC, impression share"),
    ("23", "fetch_tiktok_ads_data", "TikTok Ads API", "CPM, audience reach"),
    ("24", "fetch_linkedin_ads_data", "LinkedIn Marketing API", "CPC, audience size by title"),
]
add_table(["#", "Function", "Source", "Data"], api_ads, [0.3, 2.3, 1.8, 2.4])

doc.add_heading("Company Intelligence (4 APIs)", level=3)
api_company = [
    ("25", "fetch_company_info", "Wikipedia/Wikidata", "Company overview, HQ, industry"),
    ("26", "fetch_company_metadata", "SEC EDGAR", "Revenue, filings"),
    ("27", "fetch_sec_company_data", "SEC EDGAR", "Company financials"),
    ("28", "fetch_company_logo", "Clearbit/Logo.dev", "Company logos"),
]
add_table(["#", "Function", "Source", "Data"], api_company, [0.3, 2.3, 1.8, 2.4])

doc.add_heading("Job Market & Other (5 APIs)", level=3)
api_other = [
    ("29", "fetch_job_market", "Adzuna", "Job postings, avg salaries"),
    ("30", "fetch_jooble_data", "Jooble", "International job postings"),
    ("31", "fetch_search_trends", "Google Trends (SerpAPI)", "Search interest over time"),
    ("32", "fetch_competitor_logos", "Clearbit/Logo.dev", "Competitor logos"),
    ("33", "fetch_currency_rates", "ECB", "Exchange rates"),
]
add_table(["#", "Function", "Source", "Data"], api_other, [0.3, 2.3, 1.8, 2.4])

# C. LLM Providers
doc.add_heading("C. LLM Providers (4 via LLM Router)", level=2)
llm_data = [
    ("1", "Gemini 2.0 Flash", "gemini-2.0-flash", "Google", "15 RPM / 1,500 RPD"),
    ("2", "Groq Llama 3.3 70B", "llama-3.3-70b-versatile", "OpenAI", "30 RPM / 14,400 RPD"),
    ("3", "Cerebras Llama 3.3 70B", "llama-3.3-70b", "OpenAI", "30 RPM / 14,400 RPD"),
    ("4", "Claude (Anthropic)", "claude-sonnet-4-20250514", "Anthropic", "50 RPM / 10,000 RPD"),
]
add_table(["Priority", "Provider", "Model ID", "API Style", "Rate Limits"], llm_data, [0.6, 1.8, 2.0, 0.8, 1.6])

add_body("Strategy: Free-tier first (Gemini/Groq/Cerebras), Claude as paid fallback with circuit breakers per provider.")
doc.add_paragraph()

task_routing = [
    ("STRUCTURED (default)", "Gemini -> Groq -> Cerebras -> Claude"),
    ("CONVERSATIONAL", "Groq -> Cerebras -> Gemini -> Claude"),
    ("COMPLEX", "Groq -> Cerebras -> Gemini -> Claude"),
    ("CODE", "Gemini -> Groq -> Cerebras -> Claude"),
]
add_table(["Task Type", "Provider Priority Order"], task_routing, [2.0, 4.8])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. PRODUCT MATRIX
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Product Matrix (4 Products x 9 Data Layers)", level=1)

add_body("YES = Direct import | ORCH = Via data_orchestrator.py (lazy-loaded) | PART = Receives pre-computed data | -- = Intentionally excluded")

matrix_data = [
    ("Excel/PPT",      "YES", "YES",  "YES",  "YES",  "YES",  "YES",  "--",  "YES",  "YES"),
    ("Nova Chat",      "YES", "ORCH", "ORCH", "--",   "ORCH", "ORCH", "YES", "ORCH", "ORCH"),
    ("Slack Bot",      "YES", "ORCH", "ORCH", "--",   "ORCH", "ORCH", "YES", "ORCH", "ORCH"),
    ("PPT Generator",  "YES", "PART", "YES",  "PART", "PART", "--",   "--",  "PART", "PART"),
]
add_table(
    ["Product", "JSON\nFiles", "API\nEnrich", "Research", "Synthe-\nsizer", "Budget\nEngine", "Standar-\ndizer", "Claude\nAPI", "Trend\nEngine", "Collar\nIntel"],
    matrix_data,
    [1.1, 0.55, 0.55, 0.65, 0.6, 0.6, 0.6, 0.55, 0.6, 0.55]
)

doc.add_heading("Product Details", level=2)
product_details = [
    ("Excel/PPT", "POST /api/generate", "ZIP (Excel + PPT + HTML)", "Direct imports -- full pipeline"),
    ("Nova Chat", "POST /api/chat", "JSON (response + sources + confidence)", "Via data_orchestrator.py lazy loaders"),
    ("Slack Bot", "Slack Events API", "Slack messages", "Via nova.py -> data_orchestrator.py"),
    ("PPT Generator", "Called by app.py", "PowerPoint (.pptx)", "Receives pre-computed data from caller"),
]
add_table(["Product", "Entry Point", "Output", "Data Access Pattern"], product_details, [1.1, 1.5, 2.2, 2.0])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. DATA ORCHESTRATOR
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Data Orchestrator -- 17 Enrichment Functions", level=1)

add_body("All results cached with TTL via _api_result_cache (thread-safe). Request ID traced through all calls.")

enrich_data = [
    ("1",  "enrich_salary",               "BLS OES, Adzuna, H-1B",                  "Salary ranges, percentiles, prevailing wage"),
    ("2",  "enrich_location",             "Census, GeoNames, Teleport, DataUSA",    "Demographics, cost of living, quality of life"),
    ("3",  "enrich_market_demand",        "JOLTS, FRED, O*NET, Eurostat, ILO",      "Job openings, unemployment, skills demand"),
    ("4",  "enrich_competitive",          "Wikipedia, SEC EDGAR",                    "Company info, revenue, competitors"),
    ("5",  "enrich_employer_brand",       "Clearbit, Logo.dev",                      "Company logos, brand assets"),
    ("6",  "enrich_ad_benchmarks",        "Google, Meta, Bing, TikTok, LinkedIn",   "CPC, CPM, audience sizes, competition"),
    ("7",  "enrich_collar_intelligence",  "collar_intelligence.py",                  "Blue/white/pink collar classification"),
    ("8",  "enrich_hiring_trends",        "trend_engine.py + FRED",                  "CPC trends, seasonal patterns, momentum"),
    ("9",  "enrich_budget",               "budget_engine.py",                        "Channel allocation, CPA, hire projections"),
    ("10", "enrich_hiring_regulations",   "KB files",                                "Location-specific hiring regulations"),
    ("11", "enrich_seasonal",             "KB + trend_engine",                       "Seasonal hiring patterns by industry"),
    ("12", "enrich_campus",               "KB files",                                "Campus recruiting opportunities"),
    ("13", "enrich_events",               "KB files",                                "Hiring events, job fairs"),
    ("14", "enrich_platform_audiences",   "KB files",                                "Platform audience demographics"),
    ("15", "enrich_global_supply",        "KB + World Bank",                         "Global talent supply indicators"),
    ("16", "enrich_educational_partners", "KB files",                                "Educational institution partnerships"),
    ("17", "enrich_radio_podcasts",       "KB files",                                "Radio/podcast advertising opportunities"),
]
add_table(["#", "Function", "API Sources", "Data Produced"], enrich_data, [0.3, 2.0, 2.2, 2.3])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. TIER 2 INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Tier 2 Infrastructure Modules (5 Modules)", level=1)

tier2_data = [
    ("eval_framework.py",      "EvalSuite",            "Budget/collar/geographic/chat eval scoring"),
    ("data_contracts.py",      "validate_kb_file",     "KB JSON schema validation, enrichment output contracts"),
    ("regression_detector.py", "run_regression_check", "10 reference scenarios, baseline drift detection"),
    ("llm_router.py",         "call_llm",             "Multi-provider LLM routing with circuit breakers"),
    ("monitoring.py",         "MetricsCollector",     "Structured JSON logging, SLO monitoring, request tracing"),
]
add_table(["Module", "Key Export", "Description"], tier2_data, [1.8, 1.8, 3.2])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. NOVA CHAT TOOLS
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Nova Chat Tools (26 Tools)", level=1)

doc.add_heading("KB Direct Access (9 tools)", level=2)
kb_tools = [
    ("1", "query_global_supply",          "global_supply.json"),
    ("2", "query_channels",               "channels_db.json"),
    ("3", "query_publishers",             "joveo_publishers.json"),
    ("4", "query_knowledge_base",         "recruitment_industry_knowledge.json"),
    ("5", "query_linkedin_guidewire",     "linkedin_guidewire_data.json"),
    ("6", "query_platform_deep",          "platform_intelligence_deep.json"),
    ("7", "query_recruitment_benchmarks", "recruitment_benchmarks_deep.json"),
    ("8", "query_regional_market",        "regional_hiring_intelligence.json"),
    ("9", "query_supply_ecosystem",       "supply_ecosystem_intelligence.json"),
]
add_table(["#", "Tool", "Source"], kb_tools, [0.3, 2.5, 4.0])

doc.add_heading("KB + Intelligence (3 tools)", level=2)
kb_intel_tools = [
    ("10", "query_workforce_trends", "workforce_trends_intelligence.json"),
    ("11", "query_white_papers",     "industry_white_papers.json"),
    ("12", "suggest_smart_defaults", "All KB + orchestrator"),
]
add_table(["#", "Tool", "Source"], kb_intel_tools, [0.3, 2.5, 4.0])

doc.add_heading("Orchestrator-Backed (8 tools)", level=2)
orch_tools = [
    ("13", "query_salary_data",      "BLS + orchestrator"),
    ("14", "query_market_demand",    "JOLTS + FRED + orchestrator"),
    ("15", "query_budget_projection","budget_engine via orchestrator"),
    ("16", "query_location_profile", "Census + GeoNames + orchestrator"),
    ("17", "query_ad_platform",      "Google/Meta/LinkedIn Ads"),
    ("18", "query_employer_branding","orchestrator"),
    ("19", "query_employer_brand",   "orchestrator"),
    ("20", "query_ad_benchmarks",    "orchestrator"),
]
add_table(["#", "Tool", "Source"], orch_tools, [0.3, 2.5, 4.0])

doc.add_heading("Intelligence Engine (3 tools)", level=2)
intel_tools = [
    ("21", "query_hiring_insights", "orchestrator"),
    ("22", "query_collar_strategy", "collar_intelligence"),
    ("23", "query_market_trends",   "trend_engine"),
]
add_table(["#", "Tool", "Source"], intel_tools, [0.3, 2.5, 4.0])

doc.add_heading("v3.1 Advanced Tools (3 tools)", level=2)
v31_tools = [
    ("24", "query_role_decomposition", "collar_intelligence", "Micro1"),
    ("25", "simulate_what_if",         "budget_engine",       "Palantir"),
    ("26", "query_skills_gap",         "collar_intelligence", "Micro1"),
]
add_table(["#", "Tool", "Source", "CTO Lens"], v31_tools, [0.3, 2.2, 2.2, 2.1])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. DATA FLOW ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Data Flow Architecture", level=1)

# Stage 1
doc.add_heading("Stage 1: Input Validation & Standardization", level=2)
add_body("Entry point: app.py receives user input (industry, roles, locations, budget, client_name, competitors).")
stage1 = [
    ("standardizer.normalize_industry()", "Map free-text industry to canonical taxonomy"),
    ("standardizer.normalize_role()", "Map role to standardized title + SOC code"),
    ("collar_intelligence.classify_collar()", "Classify each role as blue/white/pink collar"),
    ("_validate_locations()", "Verify locations against US metros / country database"),
]
add_table(["Function", "Action"], stage1, [2.8, 4.0])

# Stage 2
doc.add_heading("Stage 2: Enrichment Layer (17 Functions)", level=2)
add_body("data_orchestrator.py calls 17 enrich_* functions, each pulling from 1-5 APIs. All results cached with TTL. Request ID propagated for tracing.")

stage2 = [
    ("Salary", "BLS OES, Adzuna, H-1B", "Salary ranges, percentiles"),
    ("Location", "Census, GeoNames, Teleport, DataUSA", "Demographics, cost of living"),
    ("Market Demand", "JOLTS, FRED, O*NET, ILO, Eurostat", "Job openings, skills demand"),
    ("Company Intel", "Wikipedia, SEC EDGAR", "Company overview, financials"),
    ("Brand Assets", "Clearbit, Logo.dev", "Logos"),
    ("Ad Benchmarks", "Google, Meta, Bing, TikTok, LinkedIn Ads", "CPC, CPM, audiences"),
    ("Collar Strategy", "collar_intelligence.py", "Collar classification, strategy"),
    ("Hiring Trends", "trend_engine.py + FRED", "CPC trends, seasonal, momentum"),
    ("Budget Allocation", "budget_engine.py", "Channel allocation, CPA, projections"),
]
add_table(["Enrichment Area", "API Sources", "Output"], stage2, [1.5, 2.8, 2.5])

# Stage 3
doc.add_heading("Stage 3: Cross-Source Synthesis", level=2)
add_body("data_synthesizer.py merges all 17 enrichment outputs into a unified structure with structured confidence scoring (per-field, per-source), source attribution, and conflict resolution.")

# Stage 4
doc.add_heading("Stage 4: Budget Allocation Engine", level=2)
add_body("budget_engine.py performs channel selection (collar-weighted), dynamic CPC pricing (market + seasonal + regional adjustments), seniority decomposition, channel quality scoring, budget distribution, hire projections, ROI calculation, and what-if simulation.")

# Stage 5
doc.add_heading("Stage 5: Output Generation (4 Products)", level=2)
outputs = [
    ("Excel/PPT Report", "ZIP containing: media_plan.xlsx + strategy_deck.pptx + report.html", "Full synthesized + allocated data"),
    ("Nova Chat", "JSON with response, sources, confidence, tools_used", "26 tools + LLM via llm_router"),
    ("Slack Bot", "Slack messages with formatted responses", "Same 26 tools via Slack Events API"),
    ("PPT Generator", "PowerPoint deck with exec summary, strategy, benchmarks", "Pre-computed data from app.py"),
]
add_table(["Product", "Output Format", "Data Input"], outputs, [1.3, 3.0, 2.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. AUTONOMOUS MONITORING
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Autonomous Monitoring Systems (3 Systems, 47 Probes)", level=1)

# A. Auto QC
doc.add_heading("A. Auto QC Engine (auto_qc.py)", level=2)
add_body("53 static tests + dynamic test generation + 10 self-healing strategies + weekly self-upgrade cycle.")

doc.add_heading("53 Static Tests by Category", level=3)
test_cats = [
    ("01-10", "Infrastructure", "Health endpoint, KB files, Nova import, orchestrator, Claude API, data matrix, cache, tool count, env vars, module compilation"),
    ("11-14", "Chat Quality", "Learned answers, response structure, empty message handling, ask-before-answer"),
    ("15-20", "Orchestrator & Data", "Lazy loaders, enrichment context, ad benchmarks, insights, employer brand, fallback telemetry"),
    ("21-23", "Trend Engine", "Benchmarks, seasonal patterns, data freshness"),
    ("24-28", "v3.0 Core", "Collar classification, budget engine params, ad benchmarks v3, collar intelligence, hiring trends"),
    ("29-33", "v3.0 Integration", "Nova v3 tools, PPT v3 features, synthesizer, structured confidence, data matrix v3 layers"),
    ("34-44", "v3.1 Tier 1-3", "KB data contracts, enrichment contracts, regression check, eval budget/collar, structured logging, OpenAPI, role decomposition, channel quality, dynamic CPC, what-if"),
    ("45-53", "v3.1 Coverage", "LLM router health/classify, provider availability, async jobs, API key tiers, Joveo benchmarks, audit trail, SLO compliance, tier2 modules"),
]
add_table(["Range", "Category", "Tests"], test_cats, [0.6, 1.4, 4.8])

doc.add_heading("10 Self-Healing Strategies", level=3)
heal_data = [
    ("1", "reimport_orchestrator", "Orchestrator import failure", "Reload data_orchestrator module"),
    ("2", "reimport_nova", "Nova import failure", "Reload nova module"),
    ("3", "force_matrix_recheck", "Matrix health degraded", "Force immediate data matrix re-check"),
    ("4", "cache_eviction", "Stale API cache", "Evict expired cache entries"),
    ("5", "reset_lazy_sentinels", "Lazy loader stuck in failed state", "Reset _IMPORT_FAILED sentinels"),
    ("6", "reset_llm_circuit_breakers", "LLM providers all tripped", "Reset circuit breaker states"),
    ("7", "async_job_cleanup", "Job queue overflow (>100)", "Evict completed/expired jobs"),
    ("8", "tier2_reimport", "Tier2 module import failure", "Reimport eval/contracts/regression/llm/monitoring"),
    ("9", "audit_log_persist", "Audit log exceeds max entries", "Truncate to bounded size"),
    ("10", "joveo_kb_validated", "Joveo KB file missing/corrupt", "Validate + attempt recovery"),
]
add_table(["#", "Strategy", "Trigger", "Action"], heal_data, [0.3, 1.8, 1.9, 2.8])

# B. Data Matrix Monitor
doc.add_heading("B. Data Matrix Monitor (data_matrix_monitor.py)", level=2)
add_body("47 total health probes running every 12 hours with self-healing. 60-second initial delay after startup.")

probe_summary = [
    ("Product x Layer cells", "36", "4 products x 9 data layers"),
    ("Tier2 module health", "5", "eval_framework, data_contracts, regression_detector, llm_router, monitoring"),
    ("Extended health indicators", "6", "LLM providers, async jobs, API keys, KB freshness, eval scores, regression baseline"),
    ("TOTAL", "47", ""),
]
add_table(["Category", "Count", "Details"], probe_summary, [2.0, 0.6, 4.2])

doc.add_heading("Extended Health Indicator Thresholds", level=3)
thresholds = [
    ("LLM providers", ">= 1 available", "--", "0 available"),
    ("Async job queue", "< 100 jobs", "--", ">= 100 jobs"),
    ("API keys configured", ">= 2 of 4", "1 of 4", "0 of 4"),
    ("KB file freshness", "0 stale (>90d)", "Any stale", "--"),
    ("Eval score", ">= 85%", ">= 70%", "< 70%"),
    ("Regression baseline", "< 30 days old", "< 60 days old", ">= 60 days or missing"),
]
add_table(["Indicator", "OK", "Partial", "Error"], thresholds, [1.5, 1.5, 1.8, 2.0])

# C. SLO Monitoring
doc.add_heading("C. SLO Monitoring (monitoring.py)", level=2)
slo_data = [
    ("Generate P99 latency", "< 30,000 ms", "99th percentile for /api/generate"),
    ("Chat P99 latency", "< 8,000 ms", "99th percentile for /api/chat"),
    ("Error rate", "< 1.0%", "Percentage of 5xx responses"),
    ("Availability", "> 99.5%", "Uptime percentage"),
]
add_table(["SLO", "Target", "Metric"], slo_data, [1.5, 1.5, 3.8])

add_body("Features: Structured JSON logging, X-Request-ID tracing across all 33 APIs, error budget tracking.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. API ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("8. API Endpoints", level=1)

endpoints = [
    ("POST", "/api/generate", "--", "Generate media plan (sync or async via X-Async header)"),
    ("POST", "/api/chat", "--", "Nova chat conversation"),
    ("GET", "/api/health", "--", "Basic health check"),
    ("GET", "/api/health/ready", "--", "Readiness probe"),
    ("GET", "/api/health/auto-qc", "--", "Auto QC test results"),
    ("GET", "/api/health/data-matrix", "--", "Data matrix monitoring results"),
    ("GET", "/api/health/slos", "Admin", "SLO compliance dashboard"),
    ("GET", "/api/health/eval", "Admin", "Eval framework scores"),
    ("GET", "/api/jobs/{job_id}", "--", "Poll async generation job"),
    ("GET", "/api/docs/openapi.json", "--", "OpenAPI 3.0 specification"),
    ("GET", "/docs", "--", "Swagger UI"),
    ("POST", "/api/admin/keys", "Admin", "Create/list/revoke API keys"),
    ("GET", "/api/admin/usage", "Admin", "Per-key usage dashboard"),
    ("GET", "/api/admin/audit", "Admin", "Audit trail by request_id"),
]
add_table(["Method", "Endpoint", "Auth", "Description"], endpoints, [0.5, 2.0, 0.5, 3.8])

doc.add_heading("Rate Limiting Tiers", level=2)
rate_data = [
    ("Free (per-IP)", "5", "50"),
    ("Pro (API key)", "30", "1,000"),
    ("Enterprise (API key)", "100", "10,000"),
]
add_table(["Tier", "RPM", "RPD"], rate_data, [2.0, 2.4, 2.4])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. FILE MAP
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("9. File Map", level=1)

doc.add_heading("Core Application (6 files)", level=2)
core_files = [
    ("app.py", "~9,468", "Main HTTP server, routing, generation pipeline"),
    ("api_enrichment.py", "~10,283", "33 API integrations"),
    ("data_orchestrator.py", "~2,800", "17 enrichment functions, lazy loading, caching"),
    ("data_synthesizer.py", "~1,200", "Cross-source data synthesis"),
    ("nova.py", "~3,000", "AI chat engine, 26 tools, LLM conversation"),
    ("slack_bot.py", "~600", "Slack Events API integration"),
]
add_table(["File", "Lines", "Role"], core_files, [1.8, 0.7, 4.3])

doc.add_heading("Intelligence Engines (4 files)", level=2)
intel_files = [
    ("budget_engine.py", "~2,500", "Budget allocation, quality scoring, what-if simulation"),
    ("collar_intelligence.py", "~1,500", "Collar classification, role decomposition, skills gap"),
    ("trend_engine.py", "~1,200", "CPC trends, dynamic pricing, seasonal patterns"),
    ("research.py", "~2,000", "Market data, metro data, industry intelligence"),
]
add_table(["File", "Lines", "Role"], intel_files, [1.8, 0.7, 4.3])

doc.add_heading("Output Generation (2 files)", level=2)
output_files = [
    ("ppt_generator.py", "~2,200", "PowerPoint deck generation"),
    ("standardizer.py", "~800", "Industry/role/location normalization"),
]
add_table(["File", "Lines", "Role"], output_files, [1.8, 0.7, 4.3])

doc.add_heading("Infrastructure (5 files)", level=2)
infra_files = [
    ("llm_router.py", "~748", "Multi-provider LLM routing"),
    ("monitoring.py", "~530", "Structured logging, SLO monitoring"),
    ("eval_framework.py", "~600", "AI eval suite"),
    ("data_contracts.py", "~400", "KB schema validation"),
    ("regression_detector.py", "~500", "Reference scenario drift detection"),
]
add_table(["File", "Lines", "Role"], infra_files, [1.8, 0.7, 4.3])

doc.add_heading("Monitoring & QC (2 files)", level=2)
qc_files = [
    ("auto_qc.py", "~2,100", "53 tests, 10 self-healing, weekly self-upgrade"),
    ("data_matrix_monitor.py", "~744", "47 health probes, 12-hour cycle"),
]
add_table(["File", "Lines", "Role"], qc_files, [1.8, 0.7, 4.3])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. VERIFICATION CHECKLIST
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Verification Checklist", level=1)

checks = [
    ("1", "Server health", "GET /api/health", '{"status": "ok"}'),
    ("2", "Auto QC", "GET /api/health/auto-qc", "53 tests, >90% pass"),
    ("3", "Data matrix", "GET /api/health/data-matrix", "47 probes, >90% healthy"),
    ("4", "SLO compliance", "GET /api/health/slos", "All SLOs compliant"),
    ("5", "Eval scores", "GET /api/health/eval", "Overall >85%"),
    ("6", "API docs", "GET /docs", "Swagger UI loads"),
    ("7", "Generate (sync)", "POST /api/generate", "ZIP with Excel + PPT + HTML"),
    ("8", "Generate (async)", "POST /api/generate + X-Async: true", "Job ID, poll until complete"),
    ("9", "Chat", "POST /api/chat", "JSON with response + tools_used"),
    ("10", "What-if", 'Chat: "Increase budget 30%"', "Uses simulate_what_if tool"),
    ("11", "Role decomp", 'Chat: "Break down 50 engineers"', "Uses query_role_decomposition"),
    ("12", "OpenAPI spec", "GET /api/docs/openapi.json", "Valid OpenAPI 3.0 JSON"),
]
add_table(["#", "Check", "Command/Endpoint", "Expected"], checks, [0.3, 1.2, 2.5, 2.8])


# ══════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), "v3.1_system_state.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
