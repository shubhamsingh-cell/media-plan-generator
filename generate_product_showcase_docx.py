#!/usr/bin/env python3
"""
Generate Product Showcase Document for Media Plan Generator v3.4.

A strategic product document written for CEO-level audience, showcasing
the depth, intelligence, and capability of the system built by
Shubham Singh Chandel.

This is NOT an engineering audit. It is a product capability showcase that
translates technical depth into business value and strategic impact.

Output: data/Media_Plan_Generator_Product_Showcase_v3.4.docx
"""

import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip3 install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Media_Plan_Generator_Product_Showcase_v3.4.docx")

# Brand colours
CLR_JOVEO = RGBColor(0x1A, 0x56, 0xDB)
CLR_DARK = RGBColor(0x1E, 0x1E, 0x2E)
CLR_GREY = RGBColor(0x58, 0x58, 0x6C)
CLR_LIGHT_GREY = RGBColor(0x94, 0x94, 0xA8)
CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CLR_GREEN = RGBColor(0x16, 0xA3, 0x4A)
CLR_ACCENT = RGBColor(0x7C, 0x3A, 0xED)  # Purple accent
CLR_WARM = RGBColor(0xEA, 0x58, 0x0C)    # Orange accent
CLR_ALT_ROW = RGBColor(0xF0, 0xF4, 0xFF)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

def style_header_row(row, col_count, color_hex="1A56DB"):
    set_row_height(row, 30)
    for i in range(col_count):
        cell = row.cells[i]
        set_cell_shading(cell, color_hex)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.color.rgb = CLR_WHITE
                run.font.bold = True
                run.font.size = Pt(9.5)

def add_alt_row_shading(table):
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F0F4FF")

def format_table_text(table, font_size=Pt(9)):
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = font_size

def make_table(doc, headers, rows, col_widths=None, header_color="1A56DB"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr, len(headers), header_color)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[1 + r_idx]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
    add_alt_row_shading(table)
    format_table_text(table)
    if col_widths:
        for row in table.rows:
            for c_idx, w in enumerate(col_widths):
                if c_idx < len(row.cells):
                    row.cells[c_idx].width = Inches(w)
    doc.add_paragraph()
    return table

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = CLR_DARK
    return h

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = CLR_DARK
    return p

def body_bold_start(doc, bold_text, rest_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(15)
    rb = p.add_run(bold_text)
    rb.font.bold = True
    rb.font.size = Pt(10.5)
    rb.font.color.rgb = CLR_DARK
    rr = p.add_run(rest_text)
    rr.font.size = Pt(10.5)
    rr.font.color.rgb = CLR_DARK
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def kv_line(doc, key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    rk = p.add_run(f"{key}: ")
    rk.font.bold = True
    rk.font.size = Pt(10.5)
    rk.font.color.rgb = CLR_DARK
    rv = p.add_run(str(value))
    rv.font.size = Pt(10.5)
    return p

def callout_box(doc, text):
    """Add a highlighted callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = CLR_JOVEO
    return p

def section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


# ---------------------------------------------------------------------------
# BUILD DOCUMENT
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    now = datetime.now().strftime("%B %d, %Y")

    # ======================================================================
    # COVER PAGE
    # ======================================================================
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Media Plan Generator")
    run.font.size = Pt(34)
    run.font.color.rgb = CLR_JOVEO
    run.font.bold = True

    sub = doc.add_heading(level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Product Capability Showcase")
    run.font.size = Pt(20)
    run.font.color.rgb = CLR_DARK

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run("Version 3.4  |  March 2026")
    run.font.size = Pt(14)
    run.font.color.rgb = CLR_GREY

    for _ in range(3):
        doc.add_paragraph()

    cover_lines = [
        ("Prepared by", "Shubham Singh Chandel"),
        ("Role", "Chief of Strategic Initiatives & Supply"),
        ("For", "Kshitij Jain, CEO"),
        ("Organization", "Nova AI Suite"),
        ("Classification", "Internal -- Executive Review"),
    ]
    for k, v in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        rk = p.add_run(f"{k}:  ")
        rk.font.bold = True
        rk.font.size = Pt(11)
        rk.font.color.rgb = CLR_GREY
        rv = p.add_run(v)
        rv.font.size = Pt(11)
        rv.font.color.rgb = CLR_DARK

    doc.add_page_break()

    # ======================================================================
    # TABLE OF CONTENTS
    # ======================================================================
    heading(doc, "Contents")

    toc = [
        "1.  What This Product Does",
        "2.  Why It Matters",
        "3.  What It Delivers (4 Products)",
        "4.  Nova: The AI Assistant",
        "5.  The Intelligence Layer (30 Live Data Sources)",
        "6.  How It Thinks: 12-Provider AI Engine",
        "7.  Cost Optimization: The v3.4 Milestone",
        "8.  Quality Assurance: Self-Monitoring System",
        "9.  What Has Been Fixed (62 Issues Resolved)",
        "10. Infrastructure & Deployment",
        "11. The Numbers at a Glance",
        "12. What Comes Next",
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11.5)
            run.font.color.rgb = CLR_JOVEO

    doc.add_page_break()

    # ======================================================================
    # 1. WHAT THIS PRODUCT DOES
    # ======================================================================
    heading(doc, "1. What This Product Does")

    body(doc,
        "The Media Plan Generator is an AI-powered platform that creates complete recruitment "
        "advertising media plans. Give it a client name, industry, target roles, hiring locations, "
        "and a budget -- and it returns a fully researched, data-backed media plan with channel "
        "recommendations, budget allocations, hire projections, and a ready-to-present strategy deck."
    )

    body(doc,
        "It does not guess. It pulls live data from 30 government and industry APIs -- salary "
        "benchmarks from the Bureau of Labor Statistics, job market demand from JOLTS, location "
        "demographics from the US Census, advertising costs from Google/Meta/LinkedIn, and economic "
        "indicators from the Federal Reserve, IMF, and World Bank. Every recommendation is grounded "
        "in real numbers."
    )

    body(doc,
        "It also includes Nova, a conversational AI assistant that can answer any recruitment "
        "marketing question on demand. Nova has access to 26 specialized tools covering everything "
        "from CPC benchmarks to geopolitical hiring risk, and it can simulate budget scenarios, "
        "decompose roles by seniority, and compare blue-collar vs white-collar strategies -- all "
        "through a natural conversation in the browser or on Slack."
    )

    callout_box(doc,
        "In short: this is a full-stack recruitment intelligence platform -- built from scratch, "
        "using no frameworks, with no ongoing engineering team. One person. 61,000 lines of code. "
        "25 modules. 30 live APIs. 12 AI providers. 76 automated tests. Live in production."
    )

    doc.add_page_break()

    # ======================================================================
    # 2. WHY IT MATTERS TO JOVEO
    # ======================================================================
    heading(doc, "2. Why It Matters")

    body(doc,
        "The core focus is programmatic recruitment advertising. Clients ask how "
        "to spend their hiring budgets, which channels to use, and what results to expect. "
        "Answering these questions today involves manual research, spreadsheet modeling, and "
        "tribal knowledge scattered across the team."
    )

    body(doc, "This product changes that in three ways:")

    body_bold_start(doc,
        "It standardizes media planning. ",
        "Every plan is built on the same data sources, the same allocation logic, and the "
        "same quality benchmarks. No more inconsistency between planners or regions."
    )

    body_bold_start(doc,
        "It makes institutional knowledge accessible. ",
        "The system encodes 13 curated knowledge bases (over 1 MB of structured recruitment "
        "intelligence), 42 industry research sources, and internal benchmarks into "
        "a queryable AI assistant that anyone on the team can use."
    )

    body_bold_start(doc,
        "It demonstrates what automation can look like. ",
        "This is not a proof-of-concept or a demo. It is a production system running on Render, "
        "handling real requests, with self-monitoring, self-healing, and self-upgrading quality "
        "assurance. It shows the standard of what our automation efforts can achieve."
    )

    doc.add_page_break()

    # ======================================================================
    # 3. WHAT IT DELIVERS (4 PRODUCTS)
    # ======================================================================
    heading(doc, "3. What It Delivers")

    body(doc,
        "The platform produces four distinct deliverables, all from a single input set. Each "
        "product serves a different audience and use case."
    )

    make_table(doc,
        ["Product", "Format", "Audience", "What It Contains"],
        [
            ("Media Plan Report", "Excel (.xlsx) + HTML", "Account Managers, Clients",
             "Channel allocations, budget splits, CPC/CPA benchmarks, hire projections, ROI estimates, competitor analysis"),
            ("Strategy Deck", "PowerPoint (.pptx)", "Client Presentations",
             "Executive summary, market landscape, recommended channels, budget visualization, hiring timeline"),
            ("Nova Chat", "Web API + Slack", "Internal Teams, Sales",
             "On-demand Q&A: salary benchmarks, platform comparisons, what-if simulations, hiring insights"),
            ("Slack Bot", "Slack Messages", "Internal Teams",
             "Same Nova intelligence, delivered directly in Slack channels and DMs"),
        ],
        col_widths=[1.3, 1.2, 1.4, 2.8],
    )

    heading(doc, "How a Plan Gets Built", level=2)

    body(doc, "When someone submits a plan request, the system executes a 5-stage pipeline:")

    stages = [
        ("Stage 1: Input Validation",
         "Normalizes industry names, standardizes role titles to SOC codes, "
         "classifies each role as blue-collar, white-collar, or pink-collar, "
         "and validates locations against a metro/country database."),
        ("Stage 2: Data Enrichment",
         "Fires off 17 parallel enrichment functions that pull from 30+ APIs simultaneously. "
         "Salary data from BLS. Job market demand from JOLTS and FRED. Location demographics "
         "from the Census. Ad platform costs from Google, Meta, LinkedIn, Bing, and TikTok. "
         "Company intelligence from SEC EDGAR and Wikipedia. All cached, all fault-tolerant."),
        ("Stage 3: Data Synthesis",
         "Merges all 17 enrichment outputs into a unified data structure with per-field "
         "confidence scores, source attribution, and conflict resolution. When BLS and Adzuna "
         "disagree on salary, the system picks the higher-confidence source and documents why."),
        ("Stage 4: Budget Allocation",
         "The budget engine performs channel selection (weighted by collar type), dynamic CPC "
         "pricing (adjusted for market conditions, season, and region), seniority decomposition, "
         "channel quality scoring, and ROI calculation. It can also run what-if simulations."),
        ("Stage 5: Output Generation",
         "Produces the Excel workbook with multiple sheets, the PowerPoint strategy deck with "
         "charts and recommendations, and an HTML report. All three are packaged into a ZIP "
         "file for download."),
    ]

    for stage_title, stage_desc in stages:
        body_bold_start(doc, f"{stage_title}. ", stage_desc)

    doc.add_page_break()

    # ======================================================================
    # 4. NOVA: THE AI ASSISTANT
    # ======================================================================
    heading(doc, "4. Nova: The AI Assistant")

    body(doc,
        "Nova is the conversational interface to the platform's intelligence. It is not a "
        "generic chatbot -- it has 26 specialized tools that give it direct access to live "
        "data, computed insights, and simulation capabilities. When you ask Nova a question, "
        "it decides which tools to call, fetches real data, and synthesizes an answer grounded "
        "in actual numbers."
    )

    heading(doc, "What You Can Ask Nova", level=2)

    categories = [
        ("Supply & Channels",
         "\"What does the supply look like for nurses in Texas?\" -- Pulls from "
         "global_supply.json, channels_db.json, and publisher data."),
        ("Salary & Compensation",
         "\"What's the average salary for a software engineer in San Francisco?\" -- "
         "Queries BLS OES, DataUSA, and H-1B wage data through the orchestrator."),
        ("Market Demand",
         "\"How many job openings are there for warehouse workers?\" -- Pulls JOLTS data, "
         "FRED employment indicators, and O*NET occupation outlook."),
        ("Advertising Costs",
         "\"What's the CPC for nursing roles on Indeed vs LinkedIn?\" -- Returns platform-specific "
         "CPC/CPM/CTR benchmarks by industry."),
        ("Budget Scenarios",
         "\"What if we increase the budget by 30%?\" -- Runs a live simulation through the "
         "budget engine and returns projected impact on hires, CPA, and ROI."),
        ("Hiring Strategy",
         "\"Break down 50 engineer hires by seniority level\" -- Decomposes by junior/mid/senior/lead "
         "with recommended splits and CPA multipliers."),
        ("Location Intelligence",
         "\"Compare hiring in Austin vs Denver for tech roles\" -- Pulls location demographics, "
         "cost of living, quality of life scores, and local market conditions."),
        ("Risk Assessment",
         "\"What geopolitical risks should we consider for hiring in Eastern Europe?\" -- "
         "Returns risk scores, key events, and budget adjustment recommendations."),
    ]

    for cat_title, cat_desc in categories:
        body_bold_start(doc, f"{cat_title}. ", cat_desc)

    heading(doc, "Nova's 26 Tools", level=2)

    tool_groups = [
        ("Knowledge Base Access (9 tools)",
         "Direct access to 13 curated JSON knowledge bases covering global supply, channels, "
         "publishers, platform intelligence, recruitment benchmarks, regional markets, workforce "
         "trends, white papers, and employer branding strategies."),
        ("Live Data Access (8 tools)",
         "Real-time queries to BLS, JOLTS, FRED, Census, Google/Meta/LinkedIn Ads, and the "
         "data orchestrator. Salary data, market demand, ad benchmarks, location profiles, "
         "budget projections, and employer brand intelligence."),
        ("Intelligence & Strategy (6 tools)",
         "Computed insights: hiring difficulty index, collar-based strategy comparison, CPC/CPA "
         "trend forecasting with seasonal patterns, role decomposition by seniority, what-if "
         "budget simulation, and skills gap analysis."),
        ("Advanced Analysis (3 tools)",
         "Smart defaults (auto-detect budget and channel split from partial input), geopolitical "
         "risk assessment for international hiring, and employer brand intelligence for 30+ "
         "major companies."),
    ]

    for group_title, group_desc in tool_groups:
        body_bold_start(doc, f"{group_title}. ", group_desc)

    doc.add_page_break()

    # ======================================================================
    # 5. THE INTELLIGENCE LAYER
    # ======================================================================
    heading(doc, "5. The Intelligence Layer")

    body(doc,
        "The platform's intelligence comes from three sources: 13 curated knowledge bases, "
        "30 live API integrations, and a multi-source data synthesis engine. This section "
        "details what feeds into every recommendation."
    )

    heading(doc, "13 Curated Knowledge Bases (~1 MB)", level=2)

    body(doc,
        "These are structured JSON files containing recruitment intelligence that has been "
        "researched, validated, and organized for machine consumption. They cover proprietary "
        "benchmarks, industry hiring patterns, platform pricing models, and strategic playbooks."
    )

    kb_data = [
        ("Regional Hiring Intelligence", "144 KB", "Labor market data by region"),
        ("Global Supply Data", "107 KB", "Talent supply by country/region"),
        ("Publisher Network", "96 KB", "Publisher network (Indeed, ZipRecruiter, etc.)"),
        ("Supply Ecosystem Intelligence", "90 KB", "Talent pipeline ecosystem mapping"),
        ("Platform Intelligence (Deep)", "84 KB", "CPC/CPA/CTR by platform, detailed"),
        ("Industry White Papers", "75 KB", "Strategic insights from 42 sources"),
        ("Recruitment Strategy Intelligence", "72 KB", "Playbooks by collar/tier"),
        ("Workforce Trends Intelligence", "72 KB", "Macro workforce trend data"),
        ("Recruitment Benchmarks (Deep)", "64 KB", "Conversion rates, funnel metrics"),
        ("Channels Database", "59 KB", "Channel taxonomy, pricing models"),
        ("LinkedIn Guidewire Data", "28 KB", "LinkedIn audience targeting data"),
        ("2026 Benchmarks", "11 KB", "Internal CPA/CPC by occupation"),
        ("Recruitment Industry Knowledge", "101 KB", "Industry-specific hiring patterns"),
    ]

    make_table(doc,
        ["Knowledge Base", "Size", "Content"],
        kb_data,
        col_widths=[2.5, 0.6, 3.6],
    )

    heading(doc, "30 Live API Integrations", level=2)

    body(doc,
        "Every API call has an 8-second timeout, is cached for 24 hours (in-memory + disk), "
        "fails gracefully with curated fallback data, and runs concurrently (15 parallel workers). "
        "The system never crashes because an API is down -- it always has fallback data ready."
    )

    api_groups = [
        ("Labor Market & Economics (13 APIs)", [
            "BLS OES (salary data)", "BLS QCEW (industry employment)", "BLS JOLTS (openings/hires/quits)",
            "FRED (unemployment, CPI, wages)", "FRED Employment (hourly earnings, ECI)",
            "O*NET (skills, education, outlook)", "H-1B Visa Wages (prevailing wages)",
            "CareerOneStop (DOL certifications)", "IMF (international GDP, inflation)",
            "Eurostat (EU labor stats)", "ILO/ILOSTAT (global labor data)",
            "DataUSA (occupation wages)", "World Bank (global economic indicators)",
        ]),
        ("Location & Demographics (6 APIs)", [
            "US Census ACS (population, income, education)",
            "REST Countries (population, currency, languages)",
            "GeoNames (city coordinates, timezone)",
            "Teleport (quality of life, cost of living)",
            "DataUSA Location (employment by location)",
            "Currency Rates (ECB exchange rates)",
        ]),
        ("Advertising Platforms (5 APIs)", [
            "Google Ads (CPC, impressions, competition)",
            "Meta/Facebook (CPM, audience sizing)",
            "LinkedIn Marketing (CPC, audience by title)",
            "Microsoft/Bing Ads (CPC, impression share)",
            "TikTok Marketing (CPM, audience reach)",
        ]),
        ("Company & Job Market (6 APIs)", [
            "Wikipedia REST (company descriptions)",
            "SEC EDGAR (public company filings/revenue)",
            "Clearbit (company metadata, logos)",
            "Google Favicons (company icons)",
            "Adzuna (job postings, avg salaries)",
            "Jooble (international jobs, 69 countries)",
        ]),
    ]

    for group_title, items in api_groups:
        body_bold_start(doc, group_title, "")
        for item in items:
            bullet(doc, item)

    doc.add_page_break()

    # ======================================================================
    # 6. HOW IT THINKS: 12-PROVIDER AI ENGINE
    # ======================================================================
    heading(doc, "6. How It Thinks: 12-Provider AI Engine")

    body(doc,
        "The system does not rely on a single AI provider. It has a smart routing engine "
        "that distributes requests across 12 different AI providers based on the type of "
        "question being asked. 9 of these providers are free. Only 3 are paid, and those "
        "are used only as a last resort."
    )

    body(doc,
        "When a question comes in, the router classifies it into one of 8 categories "
        "(structured data lookup, conversational advice, complex analysis, code generation, "
        "fact verification, market research, narrative writing, or batch processing) and "
        "routes it to the provider best suited for that type. If that provider is down or "
        "slow, it automatically falls through to the next one."
    )

    make_table(doc,
        ["Priority", "AI Provider", "Model", "Cost", "Strength"],
        [
            ("1", "Google Gemini", "Gemini 2.0 Flash", "Free", "Structured data, JSON, verification"),
            ("2", "Groq", "Llama 3.3 70B", "Free", "Conversational, tool calling, reasoning"),
            ("3", "Cerebras", "Llama 3.3 70B", "Free", "Redundancy (same model, different infrastructure)"),
            ("4", "Mistral", "Mistral Small", "Free", "JSON handling, multilingual"),
            ("5", "OpenRouter", "Llama 4 Maverick", "Free", "Gateway to multiple free models"),
            ("6", "xAI", "Grok 3 Mini", "Free", "Strong reasoning, 128K context window"),
            ("7", "SambaNova", "Llama 3.1 405B", "Free", "Largest open model (405 billion parameters)"),
            ("8", "NVIDIA NIM", "Llama 3.1 70B", "Free", "Hardware-optimized fast inference"),
            ("9", "Cloudflare", "Llama 3.3 70B", "Free", "Edge-distributed, low latency"),
            ("10", "OpenAI", "GPT-4o", "Paid", "Structured reasoning (fallback only)"),
            ("11", "Anthropic", "Claude Sonnet", "Paid", "Complex tool chains (last resort)"),
            ("12", "Anthropic", "Claude Opus", "Paid", "Emergency only (highest quality)"),
        ],
        col_widths=[0.6, 1.2, 1.5, 0.5, 2.9],
    )

    body(doc,
        "Each provider has its own circuit breaker -- if a provider fails 5 times in a row, "
        "it is automatically bypassed for 60 seconds. The entire routing loop has a 60-second "
        "timeout budget, so no request ever hangs indefinitely. The system is designed to always "
        "return a response, even if every AI provider is down (it falls back to a rule-based "
        "engine built on the curated knowledge bases)."
    )

    doc.add_page_break()

    # ======================================================================
    # 7. COST OPTIMIZATION: THE v3.4 MILESTONE
    # ======================================================================
    heading(doc, "7. Cost Optimization: The v3.4 Milestone")

    body(doc,
        "The most significant change in v3.4 is the migration of tool-calling from paid AI "
        "(Claude, at ~$0.003-0.005 per query) to free AI providers. Previously, any question "
        "that required looking up data -- salary benchmarks, CPC costs, market demand -- was "
        "routed to Claude because it was the only provider configured for tool calling. That "
        "meant every substantive question cost money."
    )

    body(doc,
        "In v3.4, 8 free providers now handle tool-calling natively. The system converts tool "
        "definitions between Anthropic and OpenAI formats automatically, runs multi-turn tool "
        "conversations on free providers, and only falls back to Claude if the free providers "
        "fail or return low-quality results."
    )

    make_table(doc,
        ["Query Type", "Cost Before v3.4", "Cost After v3.4", "Savings"],
        [
            ("Simple Q&A (no data)", "$0 (already free)", "$0", "--"),
            ("Data lookup (1 tool call)", "~$0.003 (Claude)", "$0 (free LLM)", "100%"),
            ("Multi-step analysis (3 tools)", "~$0.009 (Claude)", "$0 (free LLM)", "100%"),
            ("Complex tool chain (8 tools)", "~$0.024 (Claude)", "$0 free, $0.024 fallback", "~90%"),
            ("Verification / fact-check", "$0 (Gemini)", "$0 (Gemini)", "--"),
        ],
        col_widths=[2.2, 1.5, 1.7, 1.3],
    )

    callout_box(doc,
        "The practical impact: the vast majority of Nova queries now cost zero. Claude is "
        "only invoked when free providers genuinely cannot handle the request -- which, based "
        "on early results, is a small minority of queries."
    )

    doc.add_page_break()

    # ======================================================================
    # 8. QUALITY ASSURANCE
    # ======================================================================
    heading(doc, "8. Quality Assurance: Self-Monitoring System")

    body(doc,
        "The platform does not just run -- it watches itself. Three independent monitoring "
        "systems run continuously, detecting problems and fixing them automatically."
    )

    heading(doc, "Autonomous QC Engine (76 Tests)", level=2)

    body(doc,
        "Every 12 hours, the system runs a full test suite of 76 automated tests covering "
        "API endpoints, data integrity, chatbot quality, security, and performance. If a test "
        "fails, the system attempts auto-repair (cache clearing, module reloading, state reset) "
        "before alerting. Every week, it analyzes user interaction patterns and generates new "
        "test cases automatically."
    )

    kv_line(doc, "Total Tests", "76")
    kv_line(doc, "Current Pass Rate", "96.1% (73/76)")
    kv_line(doc, "Self-Healing Strategies", "10 (cache eviction, module reload, circuit reset, etc.)")
    kv_line(doc, "Test Generation", "Automatic weekly, based on production usage patterns")

    heading(doc, "Data Matrix Monitor (47 Probes)", level=2)

    body(doc,
        "47 health probes monitor every data source, every API integration, and every internal "
        "module. The monitor tracks response times, error rates, cache freshness, and data "
        "quality. When a data source degrades, the system automatically switches to fallback "
        "data and logs the event."
    )

    heading(doc, "SLO Monitoring", level=2)

    make_table(doc,
        ["SLO", "Target", "What It Measures"],
        [
            ("Plan Generation Latency", "< 30 seconds (P99)", "Time to produce a complete media plan"),
            ("Chat Response Latency", "< 8 seconds (P99)", "Time for Nova to answer a question"),
            ("Error Rate", "< 1%", "Percentage of requests returning errors"),
            ("Availability", "> 99.5%", "System uptime"),
        ],
        col_widths=[2.0, 1.5, 3.2],
    )

    doc.add_page_break()

    # ======================================================================
    # 9. WHAT HAS BEEN FIXED
    # ======================================================================
    heading(doc, "9. What Has Been Fixed")

    body(doc,
        "Across v3.3 and v3.4, a deep quality audit identified and resolved 62 issues: "
        "37 from static code analysis and 25 from stress testing under extreme conditions. "
        "Every issue has been fixed and verified."
    )

    make_table(doc,
        ["Severity", "Count", "Examples"],
        [
            ("Critical", "9", "Memory leaks, missing budget allocation in async path, "
             "thread explosion under load, CORS on errors, metrics accuracy, tool-call routing failure"),
            ("High", "11", "Lock contention, disk cache growth, N+1 database queries, "
             "thread ID propagation, sequential API calls, paid-provider guard"),
            ("Medium", "7", "Email retry storms, QC code validation, budget threshold, "
             "tool name drift, documentation sync"),
            ("Low/Info", "35", "Logging improvements, code style, documentation, "
             "edge case handling, tool call count limits"),
        ],
        col_widths=[1.0, 0.6, 5.1],
    )

    body(doc,
        "The stress test specifically evaluated behavior under 50+ concurrent requests, "
        "simultaneous API failures, memory pressure, and adversarial inputs. Every critical "
        "finding from the stress test has been resolved."
    )

    doc.add_page_break()

    # ======================================================================
    # 10. INFRASTRUCTURE & DEPLOYMENT
    # ======================================================================
    heading(doc, "10. Infrastructure & Deployment")

    body(doc,
        "The system runs on Render.com (Standard tier) with auto-deploy from GitHub. "
        "Every push to the main branch triggers a build and deployment automatically. "
        "The system is configured with 27 environment variables covering all 12 AI providers, "
        "data API keys, monitoring integrations, and authentication."
    )

    heading(doc, "Key Integrations", level=2)

    integrations = [
        ("Slack", "Bot integration for Nova chat directly in Slack channels and DMs"),
        ("Grafana + Loki", "Centralized log aggregation and dashboard visualization"),
        ("Supabase", "Persistent caching that survives deployments"),
        ("Resend", "Email alerting on critical failures"),
        ("GitHub", "Source control with auto-deploy on push"),
    ]

    make_table(doc,
        ["Integration", "Purpose"],
        integrations,
        col_widths=[1.5, 5.2],
    )

    heading(doc, "API Endpoints", level=2)

    body(doc,
        "The platform exposes 16 API endpoints covering plan generation, chatbot, health "
        "monitoring, admin controls, and API documentation. It includes a full OpenAPI 3.0 "
        "specification with interactive Swagger UI."
    )

    endpoints_summary = [
        ("Plan Generation", "/api/generate, /api/jobs/{id}", "Create and poll media plans"),
        ("AI Chat", "/api/chat, /api/nova/chat", "Nova conversational interface"),
        ("Slack", "/api/slack/events", "Slack bot webhook"),
        ("Health", "/api/health, /ready, /data-matrix, /auto-qc, /slos, /eval", "5 health endpoints"),
        ("Admin", "/api/admin/usage, /admin/keys, /admin/nova, /metrics", "Usage analytics, key management"),
        ("Documentation", "/api/docs/openapi.json, /docs", "OpenAPI spec + Swagger UI"),
    ]

    make_table(doc,
        ["Category", "Endpoints", "Purpose"],
        endpoints_summary,
        col_widths=[1.3, 3.0, 2.4],
    )

    doc.add_page_break()

    # ======================================================================
    # 11. THE NUMBERS AT A GLANCE
    # ======================================================================
    heading(doc, "11. The Numbers at a Glance")

    body(doc,
        "A summary of everything the platform comprises, as of v3.4:"
    )

    numbers = [
        ("Lines of Code", "61,219", "Across 25 Python modules"),
        ("AI Providers", "12", "9 free + 3 paid (last resort)"),
        ("Live API Integrations", "30", "Government, ad platforms, company data"),
        ("Knowledge Bases", "13", "~1 MB of curated recruitment intelligence"),
        ("Nova Tools", "26", "Specialized data access and analysis tools"),
        ("Enrichment Functions", "17", "Parallel data fetching with caching"),
        ("QC Tests", "76", "Automated, 96.1% pass rate"),
        ("Health Probes", "47", "Continuous data source monitoring"),
        ("Self-Healing Strategies", "10", "Automatic recovery from failures"),
        ("API Endpoints", "16", "REST API with OpenAPI 3.0 documentation"),
        ("Task Classifications", "8", "Intelligent routing by question type"),
        ("Environment Variables", "27", "All configured on Render"),
        ("Bug Fixes (v3.3-3.4)", "62", "9 critical, 11 high, 7 medium"),
        ("Output Products", "4", "Excel, PowerPoint, Web Chat, Slack Bot"),
        ("Git Commits", "70+", "Full version history since inception"),
        ("External Dependencies", "2", "openpyxl + python-pptx (everything else is stdlib)"),
    ]

    make_table(doc,
        ["Metric", "Value", "Detail"],
        numbers,
        col_widths=[1.8, 0.8, 4.1],
    )

    callout_box(doc,
        "Built entirely by one person, using no web frameworks, with only 2 external "
        "dependencies (Excel and PowerPoint libraries). The HTTP server, routing, caching, "
        "rate limiting, monitoring, circuit breakers, and all AI integrations are pure "
        "Python standard library."
    )

    doc.add_page_break()

    # ======================================================================
    # 12. WHAT COMES NEXT
    # ======================================================================
    heading(doc, "12. What Comes Next")

    heading(doc, "Near-Term (This Month)", level=2)

    bullet(doc,
        "Monitor free LLM tool-calling performance in production and tune provider "
        "priority order based on observed success rates and latency."
    )
    bullet(doc,
        "Add structured logging to track exactly how much the cost optimization is "
        "saving -- every query tagged with which provider handled it and at what cost."
    )
    bullet(doc,
        "Build a response quality comparison framework to validate that free providers "
        "are matching Claude's answer quality on tool-calling queries."
    )

    heading(doc, "Medium-Term (1-3 Months)", level=2)

    bullet(doc,
        "Streaming responses for Nova to reduce perceived wait time in conversations."
    )
    bullet(doc,
        "Cost and usage dashboard showing real-time spend across all 12 providers."
    )
    bullet(doc,
        "A/B testing framework to systematically compare provider quality."
    )
    bullet(doc,
        "Expand dynamic test generation to analyze production error logs with AI."
    )

    heading(doc, "Long-Term (3-6 Months)", level=2)

    bullet(doc,
        "Fine-tune a smaller model on Nova's specific tool-calling patterns for even "
        "faster inference at zero cost."
    )
    bullet(doc,
        "Multi-language support for the chatbot, leveraging Mistral's multilingual strengths."
    )
    bullet(doc,
        "Self-serve analytics dashboard for clients to explore their own media plan data."
    )

    # ======================================================================
    # CLOSING
    # ======================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("---")
    run.font.color.rgb = CLR_GREY
    run.font.size = Pt(14)

    doc.add_paragraph()

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_after = Pt(4)
    run = closing.add_run("Media Plan Generator v3.4")
    run.font.size = Pt(12)
    run.font.color.rgb = CLR_JOVEO
    run.font.bold = True

    closing2 = doc.add_paragraph()
    closing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing2.paragraph_format.space_after = Pt(4)
    run = closing2.add_run("Built by Shubham Singh Chandel")
    run.font.size = Pt(11)
    run.font.color.rgb = CLR_DARK

    closing3 = doc.add_paragraph()
    closing3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing3.add_run("Creator  |  Nova AI Suite")
    run.font.size = Pt(10)
    run.font.color.rgb = CLR_GREY

    closing4 = doc.add_paragraph()
    closing4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing4.add_run(now)
    run.font.size = Pt(10)
    run.font.color.rgb = CLR_GREY

    # ======================================================================
    # SAVE
    # ======================================================================
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Report saved to: {OUTPUT_FILE}")
    print(f"File size: {os.path.getsize(OUTPUT_FILE):,} bytes")
    return OUTPUT_FILE


if __name__ == "__main__":
    build()
