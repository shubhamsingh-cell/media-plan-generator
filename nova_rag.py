"""
Nova RAG (Retrieval-Augmented Generation) Module

Handles document extraction, embedding, and retrieval for context in conversations.
Supports PDF, DOCX, TXT, and XLSX files.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

# Try to import document extraction libraries
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl

    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False


# ---------------------------------------------------------------------------
# Document Extraction
# ---------------------------------------------------------------------------


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """Extract text from PDF file.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text, or None on error
    """
    if not PDF_AVAILABLE:
        logger.warning("PyPDF2 not available; cannot extract PDF")
        return None

    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text.strip() if text else None
    except Exception as e:
        logger.error("Error extracting PDF: %s", e, exc_info=True)
        return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """Extract text from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted text, or None on error
    """
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not available; cannot extract DOCX")
        return None

    try:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
        return text.strip() if text else None
    except Exception as e:
        logger.error("Error extracting DOCX: %s", e, exc_info=True)
        return None


def extract_text_from_xlsx(file_path: str) -> Optional[str]:
    """Extract text from XLSX file.

    Args:
        file_path: Path to XLSX file

    Returns:
        Extracted text (tab-separated), or None on error
    """
    if not XLSX_AVAILABLE:
        logger.warning("openpyxl not available; cannot extract XLSX")
        return None

    try:
        wb = openpyxl.load_workbook(file_path)
        text = ""
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text += f"\n=== Sheet: {sheet} ===\n"
            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join(str(v) if v is not None else "" for v in row)
                if row_text.strip():
                    text += row_text + "\n"
        return text.strip() if text else None
    except Exception as e:
        logger.error("Error extracting XLSX: %s", e, exc_info=True)
        return None


def extract_text_from_file(file_path: str, file_type: str) -> Optional[str]:
    """Extract text from any supported file type.

    Args:
        file_path: Path to file
        file_type: 'pdf', 'docx', 'txt', 'xlsx'

    Returns:
        Extracted text, or None on error
    """
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type == "xlsx":
        return extract_text_from_xlsx(file_path)
    elif file_type == "txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read().strip() or None
        except Exception as e:
            logger.error("Error reading text file: %s", e, exc_info=True)
            return None
    else:
        logger.warning("Unsupported file type: %s", file_type)
        return None


# ---------------------------------------------------------------------------
# Text Processing for Embeddings
# ---------------------------------------------------------------------------


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks for embedding.

    Args:
        text: Text to chunk
        chunk_size: Characters per chunk
        overlap: Character overlap between chunks

    Returns:
        List of text chunks
    """
    if not text:
        return []

    chunks = []
    step = chunk_size - overlap
    for i in range(0, len(text), step):
        chunk = text[i : i + chunk_size]
        if chunk.strip():
            chunks.append(chunk.strip())
    return chunks


def clean_text(text: str) -> str:
    """Clean and normalize text for embedding.

    Args:
        text: Text to clean

    Returns:
        Cleaned text
    """
    # Remove extra whitespace
    text = re.sub(r"\s+", " ", text)

    # Remove control characters
    text = "".join(ch for ch in text if ord(ch) >= 32 or ch in "\n\t")

    # Remove URLs (optional: set to preserve for some use cases)
    text = re.sub(r"http[s]?://\S+", "[URL]", text)

    return text.strip()


def extract_key_phrases(text: str, max_phrases: int = 10) -> List[str]:
    """Extract key phrases from text (simple heuristic).

    Args:
        text: Text to extract from
        max_phrases: Max phrases to return

    Returns:
        List of key phrases
    """
    # Split by common delimiters and find noun phrases (simplified)
    words = text.split()
    phrases = []

    # Simple heuristic: look for capitalized sequences
    current_phrase = []
    for word in words:
        if word and word[0].isupper() and len(word) > 2:
            current_phrase.append(word)
        elif current_phrase:
            if len(current_phrase) >= 2:
                phrase = " ".join(current_phrase)
                if phrase not in phrases:
                    phrases.append(phrase)
            current_phrase = []

    if current_phrase and len(current_phrase) >= 2:
        phrase = " ".join(current_phrase)
        if phrase not in phrases:
            phrases.append(phrase)

    return phrases[:max_phrases]


# ---------------------------------------------------------------------------
# Simple Vector Similarity (Cosine) - No External Library
# ---------------------------------------------------------------------------


def simple_embedding(text: str) -> List[float]:
    """Generate a simple embedding using TF-IDF-like scoring.

    This is a basic implementation for demo purposes.
    For production, use proper embedding models (OpenAI, Hugging Face, etc.)

    Args:
        text: Text to embed

    Returns:
        Embedding vector (simple bag-of-words + position weighting)
    """
    # Tokenize
    words = text.lower().split()
    if not words:
        return [0.0] * 100  # Default zero vector

    # Simple embedding: word frequency + position weighting
    word_freq = {}
    for i, word in enumerate(words):
        # Clean word
        word = re.sub(r"[^\w]", "", word)
        if len(word) > 2:  # Skip short words
            position_weight = 1.0 + (i / len(words))  # Earlier words weighted higher
            word_freq[word] = word_freq.get(word, 0) + position_weight

    # Create feature vector (top 100 terms)
    top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:100]
    embedding = [freq for _, freq in top_words] + [0.0] * (100 - len(top_words))

    # Normalize
    norm = sum(x * x for x in embedding) ** 0.5
    if norm > 0:
        embedding = [x / norm for x in embedding]

    return embedding


def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    """Calculate cosine similarity between two vectors.

    Args:
        vec1: First vector
        vec2: Second vector

    Returns:
        Similarity score (0-1)
    """
    if not vec1 or not vec2 or len(vec1) != len(vec2):
        return 0.0

    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    norm1 = sum(x * x for x in vec1) ** 0.5
    norm2 = sum(x * x for x in vec2) ** 0.5

    if norm1 == 0 or norm2 == 0:
        return 0.0

    return dot_product / (norm1 * norm2)


# ---------------------------------------------------------------------------
# RAG Retrieval
# ---------------------------------------------------------------------------


def retrieve_relevant_chunks(
    query: str,
    documents: List[Dict[str, Any]],
    top_k: int = 3,
) -> List[Tuple[str, float]]:
    """Retrieve relevant document chunks for a query.

    Args:
        query: User query
        documents: List of document dicts with 'content_text'
        top_k: Number of top results to return

    Returns:
        List of (chunk_text, similarity_score) tuples
    """
    if not documents or not query:
        return []

    try:
        # Create query embedding
        query_embedding = simple_embedding(query)

        # Score all document chunks
        scored_chunks = []
        for doc in documents:
            content = doc.get("content_text") or ""
            if not content:
                continue

            # Chunk the document
            chunks = chunk_text(content, chunk_size=500, overlap=100)

            for chunk in chunks:
                chunk_embedding = simple_embedding(chunk)
                similarity = cosine_similarity(query_embedding, chunk_embedding)
                if similarity > 0.1:  # Filter low-similarity chunks
                    scored_chunks.append((chunk, similarity))

        # Sort by similarity and return top-k
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        return scored_chunks[:top_k]

    except Exception as e:
        logger.error("Error retrieving chunks: %s", e, exc_info=True)
        return []


def format_context_for_prompt(chunks: List[Tuple[str, float]]) -> str:
    """Format retrieved chunks into a context string for the LLM prompt.

    Args:
        chunks: List of (chunk_text, similarity_score) tuples

    Returns:
        Formatted context string
    """
    if not chunks:
        return ""

    context = "## Document Context:\n\n"
    for i, (chunk, score) in enumerate(chunks, 1):
        context += f"**Reference {i}** (relevance: {score:.1%})\n"
        context += f"{chunk[:500]}...\n\n" if len(chunk) > 500 else f"{chunk}\n\n"

    return context


# ---------------------------------------------------------------------------
# Document Metadata Extraction
# ---------------------------------------------------------------------------


def extract_document_metadata(file_path: str, file_type: str) -> Dict[str, Any]:
    """Extract metadata from document.

    Args:
        file_path: Path to file
        file_type: 'pdf', 'docx', 'txt', 'xlsx'

    Returns:
        Metadata dict
    """
    metadata = {
        "file_type": file_type,
        "file_path": file_path,
    }

    try:
        if file_type == "pdf" and PDF_AVAILABLE:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                metadata["page_count"] = len(reader.pages)
                if reader.metadata:
                    metadata["title"] = reader.metadata.get("/Title")
                    metadata["author"] = reader.metadata.get("/Author")

        elif file_type == "docx" and DOCX_AVAILABLE:
            doc = Document(file_path)
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            if doc.core_properties:
                metadata["title"] = doc.core_properties.title
                metadata["author"] = doc.core_properties.author

        elif file_type == "xlsx" and XLSX_AVAILABLE:
            wb = openpyxl.load_workbook(file_path)
            metadata["sheet_count"] = len(wb.sheetnames)
            metadata["sheets"] = wb.sheetnames

    except Exception as e:
        logger.warning("Could not extract metadata: %s", e)

    return metadata
