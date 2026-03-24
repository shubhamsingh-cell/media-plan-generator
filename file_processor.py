"""
file_processor.py  --  Text extraction from uploaded documents
=============================================================
Handles PDF, DOCX, TXT, VTT/SRT transcript files, and CSV/Excel
historical performance data.

All functions accept raw bytes and return extracted text or
structured data. No filesystem I/O required.
"""

import io
import re
import csv
import json
import logging
import base64
from typing import List, Dict, Any, Optional, Tuple

logger = logging.getLogger("file_processor")

# ─────────────────────────────────────────────────────────────
# Text extraction from documents
# ─────────────────────────────────────────────────────────────


def extract_text_from_base64(b64_data: str, filename: str, file_type: str = "") -> str:
    """Extract text content from a base64-encoded file.

    Parameters
    ----------
    b64_data : str
        Base64-encoded file content.
    filename : str
        Original filename (used to determine format).
    file_type : str
        MIME type hint (optional).

    Returns
    -------
    str
        Extracted text content, or empty string on failure.
    """
    try:
        raw_bytes = base64.b64decode(b64_data)
    except Exception as e:
        logger.warning("Failed to decode base64 for %s: %s", filename, e)
        return ""

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "txt" or file_type == "text/plain":
        return _extract_txt(raw_bytes, filename)
    elif ext == "pdf" or file_type == "application/pdf":
        return _extract_pdf(raw_bytes, filename)
    elif ext in ("docx", "doc") or "wordprocessingml" in file_type:
        return _extract_docx(raw_bytes, filename)
    elif ext in ("vtt", "srt"):
        return _extract_subtitle(raw_bytes, filename)
    elif ext == "csv" or file_type == "text/csv":
        return _extract_txt(raw_bytes, filename)  # CSV as text
    elif ext in ("xlsx", "xls"):
        return _extract_excel_as_text(raw_bytes, filename)
    else:
        logger.warning("Unsupported file type: %s (%s)", filename, file_type)
        return ""


def _extract_txt(raw_bytes: bytes, filename: str) -> str:
    """Extract text from plain text / CSV files."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return raw_bytes.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    logger.warning("Could not decode text file: %s", filename)
    return ""


def _extract_pdf(raw_bytes: bytes, filename: str) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(raw_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        result = "\n\n".join(pages)
        logger.info(
            "Extracted %d chars from PDF %s (%d pages)",
            len(result),
            filename,
            len(reader.pages),
        )
        return result
    except ImportError:
        logger.warning(
            "PyPDF2 not installed -- cannot extract PDF text from %s", filename
        )
        return f"[PDF file: {filename} -- install PyPDF2 for text extraction]"
    except Exception as e:
        logger.warning("PDF extraction failed for %s: %s", filename, e)
        return ""


def _extract_docx(raw_bytes: bytes, filename: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(raw_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        result = "\n".join(paragraphs)
        logger.info("Extracted %d chars from DOCX %s", len(result), filename)
        return result
    except ImportError:
        logger.warning(
            "python-docx not installed -- cannot extract DOCX text from %s", filename
        )
        return f"[DOCX file: {filename} -- install python-docx for text extraction]"
    except Exception as e:
        logger.warning("DOCX extraction failed for %s: %s", filename, e)
        return ""


def _extract_subtitle(raw_bytes: bytes, filename: str) -> str:
    """Extract text from VTT/SRT subtitle/transcript files."""
    text = _extract_txt(raw_bytes, filename)
    if not text:
        return ""
    # Strip VTT/SRT timing lines and headers
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        # Skip empty, numeric-only (SRT sequence), timing lines, and WEBVTT header
        if not line:
            continue
        if line.isdigit():
            continue
        if re.match(r"^\d{2}:\d{2}", line):
            continue
        if line.startswith("WEBVTT") or line.startswith("NOTE"):
            continue
        lines.append(line)
    return "\n".join(lines)


def _extract_excel_as_text(raw_bytes: bytes, filename: str) -> str:
    """Extract text representation from Excel files."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c for c in cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
        wb.close()
        result = "\n\n".join(parts)
        logger.info("Extracted %d chars from Excel %s", len(result), filename)
        return result
    except Exception as e:
        logger.warning("Excel extraction failed for %s: %s", filename, e)
        return ""


# ─────────────────────────────────────────────────────────────
# Batch extraction for multiple uploaded files
# ─────────────────────────────────────────────────────────────


def extract_all_texts(files: List[Dict[str, str]]) -> str:
    """Extract and concatenate text from multiple uploaded files.

    Parameters
    ----------
    files : list of dict
        Each dict has keys: name, type, data (base64).

    Returns
    -------
    str
        Combined extracted text, separated by file headers.
    """
    if not files:
        return ""
    parts = []
    for f in files:
        text = extract_text_from_base64(
            f.get("data") or "", f.get("name", "unknown"), f.get("type") or ""
        )
        if text.strip():
            parts.append(f"=== From: {f.get('name', 'unknown')} ===\n{text.strip()}")
    return "\n\n".join(parts)


# ─────────────────────────────────────────────────────────────
# Historical performance data parsing
# ─────────────────────────────────────────────────────────────


def parse_historical_data(files: List[Dict[str, str]]) -> Optional[Dict[str, Any]]:
    """Parse historical campaign performance data from CSV/Excel uploads.

    Expected columns (flexible matching):
    - Platform/Channel/Source/Publisher
    - Spend/Budget/Cost/Investment
    - Applications/Applies/Clicks
    - Hires/Conversions
    - CPA/Cost Per Application
    - CPC/Cost Per Click
    - CPH/Cost Per Hire

    Parameters
    ----------
    files : list of dict
        Uploaded files with base64 data.

    Returns
    -------
    dict or None
        Structured historical data with platform breakdown, or None if
        no valid data found.
    """
    if not files:
        return None

    all_rows: List[Dict[str, Any]] = []

    for f in files:
        try:
            raw_bytes = base64.b64decode(f.get("data") or "")
        except Exception:
            continue

        ext = (f.get("name") or "").rsplit(".", 1)[-1].lower()

        if ext == "csv":
            rows = _parse_csv_historical(raw_bytes, f.get("name") or "")
            all_rows.extend(rows)
        elif ext in ("xlsx", "xls"):
            rows = _parse_excel_historical(raw_bytes, f.get("name") or "")
            all_rows.extend(rows)

    if not all_rows:
        return None

    # Aggregate by platform
    platforms: Dict[str, Dict[str, float]] = {}
    total_spend = 0.0
    total_hires = 0
    total_applications = 0

    for row in all_rows:
        platform = row.get("platform", "Unknown")
        if platform not in platforms:
            platforms[platform] = {
                "spend": 0.0,
                "applications": 0,
                "hires": 0,
                "clicks": 0,
                "cpa": 0.0,
                "cpc": 0.0,
                "cph": 0.0,
            }
        p = platforms[platform]
        p["spend"] += row.get("spend", 0.0)
        p["applications"] += row.get("applications") or 0
        p["hires"] += row.get("hires") or 0
        p["clicks"] += row.get("clicks") or 0
        total_spend += row.get("spend", 0.0)
        total_hires += row.get("hires") or 0
        total_applications += row.get("applications") or 0

    # Calculate derived metrics per platform
    for p_data in platforms.values():
        if p_data["applications"] > 0:
            p_data["cpa"] = round(p_data["spend"] / p_data["applications"], 2)
        if p_data["clicks"] > 0:
            p_data["cpc"] = round(p_data["spend"] / p_data["clicks"], 2)
        if p_data["hires"] > 0:
            p_data["cph"] = round(p_data["spend"] / p_data["hires"], 2)

    result = {
        "platforms": platforms,
        "total_spend": round(total_spend, 2),
        "total_hires": total_hires,
        "total_applications": total_applications,
        "total_cpa": (
            round(total_spend / total_applications, 2)
            if total_applications > 0
            else 0.0
        ),
        "total_cph": round(total_spend / total_hires, 2) if total_hires > 0 else 0.0,
        "platform_count": len(platforms),
        "source_files": [f.get("name", "unknown") for f in files],
    }
    logger.info(
        "Parsed historical data: %d platforms, $%.0f total spend from %d files",
        len(platforms),
        total_spend,
        len(files),
    )
    return result


# Column name matching patterns
_COL_PATTERNS = {
    "platform": re.compile(
        r"platform|channel|source|publisher|board|vendor|site", re.I
    ),
    "spend": re.compile(
        r"spend|budget|cost|investment|total.?spend|media.?spend", re.I
    ),
    "applications": re.compile(
        r"appli|applies|leads|conversions|responses|submis", re.I
    ),
    "hires": re.compile(r"hires?|hired|offers?|placements?|starts?", re.I),
    "clicks": re.compile(r"clicks?|visits?|traffic", re.I),
    "cpa": re.compile(r"cpa|cost.?per.?app|cost.?per.?lead", re.I),
    "cpc": re.compile(r"cpc|cost.?per.?click", re.I),
    "cph": re.compile(r"cph|cost.?per.?hire", re.I),
}


def _match_columns(headers: List[str]) -> Dict[str, int]:
    """Match header row to known column names. Returns {field: col_index}."""
    mapping = {}
    for idx, header in enumerate(headers):
        header_clean = str(header).strip()
        if not header_clean:
            continue
        for field, pattern in _COL_PATTERNS.items():
            if pattern.search(header_clean) and field not in mapping:
                mapping[field] = idx
                break
    return mapping


def _parse_number(val: Any) -> float:
    """Parse a numeric value from a cell, handling currency/commas."""
    if isinstance(val, (int, float)):
        return float(val)
    if not val:
        return 0.0
    s = str(val).strip().replace(",", "").replace("$", "").replace("%", "")
    try:
        return float(s)
    except (ValueError, TypeError):
        return 0.0


def _parse_csv_historical(raw_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """Parse CSV file into historical performance rows."""
    text = _extract_txt(raw_bytes, filename)
    if not text:
        return []

    rows = []
    reader = csv.reader(io.StringIO(text))
    headers = None
    col_map = None

    for line in reader:
        if not any(cell.strip() for cell in line):
            continue
        if headers is None:
            headers = line
            col_map = _match_columns(headers)
            if "platform" not in col_map:
                # Try next non-empty row as header
                headers = None
                col_map = None
                continue
            continue

        if col_map is None:
            continue

        row_data = {}
        platform_idx = col_map.get("platform")
        if platform_idx is not None and platform_idx < len(line):
            row_data["platform"] = str(line[platform_idx]).strip()
        else:
            continue

        if not row_data["platform"]:
            continue

        for field in ("spend", "applications", "hires", "clicks", "cpa", "cpc", "cph"):
            idx = col_map.get(field)
            if idx is not None and idx < len(line):
                row_data[field] = _parse_number(line[idx])
            else:
                row_data[field] = 0.0

        # Convert applications/hires to int
        for int_field in ("applications", "hires", "clicks"):
            row_data[int_field] = int(row_data[int_field])

        rows.append(row_data)

    logger.info("Parsed %d rows from CSV %s", len(rows), filename)
    return rows


def _parse_excel_historical(raw_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """Parse Excel file into historical performance rows."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.warning("openpyxl not available for Excel parsing")
        return []

    try:
        wb = load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
    except Exception as e:
        logger.warning("Failed to open Excel %s: %s", filename, e)
        return []

    rows = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        headers = None
        col_map = None

        for excel_row in ws.iter_rows(values_only=True):
            cells = list(excel_row)
            if not any(c for c in cells if c is not None):
                continue

            if headers is None:
                headers = [str(c) if c else "" for c in cells]
                col_map = _match_columns(headers)
                if "platform" not in col_map:
                    headers = None
                    col_map = None
                    continue
                continue

            if col_map is None:
                continue

            row_data = {}
            platform_idx = col_map.get("platform")
            if platform_idx is not None and platform_idx < len(cells):
                val = cells[platform_idx]
                row_data["platform"] = str(val).strip() if val else ""
            else:
                continue

            if not row_data["platform"]:
                continue

            for field in (
                "spend",
                "applications",
                "hires",
                "clicks",
                "cpa",
                "cpc",
                "cph",
            ):
                idx = col_map.get(field)
                if idx is not None and idx < len(cells):
                    row_data[field] = _parse_number(cells[idx])
                else:
                    row_data[field] = 0.0

            for int_field in ("applications", "hires", "clicks"):
                row_data[int_field] = int(row_data[int_field])

            rows.append(row_data)

    wb.close()
    logger.info("Parsed %d rows from Excel %s", len(rows), filename)
    return rows


# ─────────────────────────────────────────────────────────────
# Transcript summarization helper
# ─────────────────────────────────────────────────────────────


def summarize_transcript_context(transcript_text: str, max_chars: int = 3000) -> str:
    """Extract key context from a sales call transcript.

    Looks for patterns indicating:
    - Hiring goals and volumes
    - Pain points with current vendors/approach
    - Budget signals and constraints
    - Timeline expectations
    - Decision criteria
    - Specific role/location mentions

    Parameters
    ----------
    transcript_text : str
        Raw transcript text.
    max_chars : int
        Maximum output length.

    Returns
    -------
    str
        Structured summary of key context extracted from transcript.
    """
    if not transcript_text or len(transcript_text.strip()) < 50:
        return ""

    text = transcript_text.strip()

    # Pattern-based extraction (no LLM needed for basic extraction)
    sections = []

    # Hiring volume signals
    hire_patterns = re.findall(
        r"(?:hiring|recruit|fill|open|need|looking for|target)\s*(?:about|around|approximately|roughly)?\s*(\d[\d,]*)\s*(?:\+?\s*)?(?:positions?|roles?|people|hires?|openings?|headcount)?",
        text,
        re.I,
    )
    if hire_patterns:
        sections.append(
            f"Hiring signals: {', '.join(set(hire_patterns[:5]))} positions mentioned"
        )

    # Budget signals
    budget_patterns = re.findall(
        r"\$\s*[\d,.]+\s*(?:[KkMmBb](?:illion)?)?(?:\s*(?:per|a|each)\s+(?:year|month|quarter))?",
        text,
    )
    if budget_patterns:
        sections.append(f"Budget signals: {', '.join(set(budget_patterns[:5]))}")

    # Pain points (common keywords)
    pain_keywords = re.findall(
        r"(?:struggling|challenge|problem|issue|difficult|expensive|slow|poor|frustrat|waste|inefficien|not working|underperform)",
        text,
        re.I,
    )
    if pain_keywords:
        sections.append(f"Pain points mentioned: {len(set(pain_keywords))} indicators")

    # Platform mentions
    known_platforms = [
        "Indeed",
        "LinkedIn",
        "ZipRecruiter",
        "Glassdoor",
        "CareerBuilder",
        "Monster",
        "Dice",
        "SimplyHired",
        "Handshake",
        "Hired",
        "Google Jobs",
        "Facebook",
        "Instagram",
        "TikTok",
        "Programmatic",
        "Appcast",
        "Pandologic",
        "Recruitics",
        "Talroo",
        "Neuvoo",
    ]
    mentioned = [p for p in known_platforms if p.lower() in text.lower()]
    if mentioned:
        sections.append(f"Platforms discussed: {', '.join(mentioned)}")

    # Timeline signals
    timeline_patterns = re.findall(
        r"(?:(?:by|before|within|next|starting)\s+)?(?:Q[1-4]|January|February|March|April|May|June|July|August|September|October|November|December|\d+\s+(?:months?|weeks?|years?|quarters?))",
        text,
        re.I,
    )
    if timeline_patterns:
        sections.append(f"Timeline references: {', '.join(set(timeline_patterns[:5]))}")

    # Location mentions (basic)
    location_patterns = re.findall(
        r"(?:in|across|throughout|covering)\s+([\w\s,]+?)(?:\.|,\s*(?:and|we|they|the)|$)",
        text[:5000],
        re.I,
    )
    if location_patterns:
        locs = [l.strip() for l in location_patterns[:3] if len(l.strip()) < 80]
        if locs:
            sections.append(f"Locations mentioned: {'; '.join(locs)}")

    if not sections:
        # Fallback: just truncate the transcript
        return f"[Call transcript provided -- {len(text)} chars]\n{text[:max_chars]}"

    summary = "KEY CONTEXT FROM SALES CALL:\n" + "\n".join(f"- {s}" for s in sections)

    # Append truncated raw text for additional context
    remaining = max_chars - len(summary) - 50
    if remaining > 200:
        summary += f"\n\nRAW TRANSCRIPT EXCERPT:\n{text[:remaining]}"

    return summary[:max_chars]
