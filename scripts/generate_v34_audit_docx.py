#!/usr/bin/env python3
"""
Generate v3.4 System Audit Report DOCX for the Media Plan Generator.

Produces a comprehensive audit document covering:
  - Executive Summary
  - Architecture Overview
  - v3.4 Free LLM Tool Calling Migration
  - 12-Provider LLM Routing Matrix
  - Bug Fixes (CRITICAL through LOW)
  - 30 API Integrations
  - 26 Nova Chatbot Tools
  - QC Test Results (76 tests)
  - Deployment & Infrastructure
  - Security & Safeguards
  - Recommendations

Output: data/System_Audit_Report_v3.4.docx
"""

import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip3 install python-docx")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "System_Audit_Report_v3.4.docx")

# Colours (RGB)
CLR_JOVEO_BLUE = RGBColor(0x1A, 0x56, 0xDB)
CLR_DARK = RGBColor(0x1E, 0x1E, 0x2E)
CLR_GREY = RGBColor(0x58, 0x58, 0x6C)
CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CLR_CRITICAL = RGBColor(0xDC, 0x26, 0x26)
CLR_HIGH = RGBColor(0xEA, 0x58, 0x0C)
CLR_MEDIUM = RGBColor(0xCA, 0x8A, 0x04)
CLR_LOW = RGBColor(0x16, 0xA3, 0x4A)
CLR_GREEN = RGBColor(0x16, 0xA3, 0x4A)
CLR_HEADER_BG = RGBColor(0x1A, 0x56, 0xDB)
CLR_ALT_ROW = RGBColor(0xF0, 0xF4, 0xFF)
CLR_BORDER = RGBColor(0xD1, 0xD5, 0xDB)

SEV_COLORS = {
    "CRITICAL": CLR_CRITICAL,
    "HIGH": CLR_HIGH,
    "MEDIUM": CLR_MEDIUM,
    "LOW": CLR_LOW,
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def style_header_row(row, col_count):
    set_row_height(row, 28)
    for i in range(col_count):
        cell = row.cells[i]
        set_cell_shading(cell, "1A56DB")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.color.rgb = CLR_WHITE
                run.font.bold = True
                run.font.size = Pt(9)


def add_alt_row_shading(table):
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F0F4FF")


def format_table_text(table, font_size=Pt(8.5)):
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = font_size


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
    style_header_row(hdr, len(headers))

    for r_idx, row_data in enumerate(rows):
        row = table.rows[1 + r_idx]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)

    add_alt_row_shading(table)
    format_table_text(table)

    if col_widths:
        for r_idx, row in enumerate(table.rows):
            for c_idx, w in enumerate(col_widths):
                if c_idx < len(row.cells):
                    row.cells[c_idx].width = Inches(w)

    doc.add_paragraph()
    return table


def add_heading_num(doc, number, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(f"{number}. {text}")
    run.font.color.rgb = CLR_DARK
    return heading


def add_kv(doc, key, value, bold_value=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run_k = p.add_run(f"{key}: ")
    run_k.font.bold = True
    run_k.font.size = Pt(10)
    run_k.font.color.rgb = CLR_DARK
    run_v = p.add_run(str(value))
    run_v.font.size = Pt(10)
    if bold_value:
        run_v.font.bold = True
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = CLR_DARK
    return p


def add_severity_badge(paragraph, severity):
    run = paragraph.add_run(f"  [{severity}]")
    run.font.color.rgb = SEV_COLORS.get(severity, CLR_GREY)
    run.font.bold = True
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# REPORT BUILDER
# ---------------------------------------------------------------------------


def build_report():
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    now = datetime.now().strftime("%B %d, %Y")

    # ====================================================================
    # TITLE PAGE
    # ====================================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Media Plan Generator")
    run.font.size = Pt(32)
    run.font.color.rgb = CLR_JOVEO_BLUE
    run.font.bold = True

    subtitle = doc.add_heading(level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("System Audit Report v3.4")
    run.font.size = Pt(22)
    run.font.color.rgb = CLR_DARK

    for _ in range(2):
        doc.add_paragraph()

    meta_items = [
        ("Version", "3.4.0"),
        ("Date", now),
        ("Platform", "Render.com (Standard Tier)"),
        ("URL", "https://media-plan-generator.onrender.com"),
        ("Repository", "github.com/shubham-chandel/media-plan-generator"),
        ("Codebase", "25 Python modules, 61,219 lines"),
        ("QC Status", "73/76 tests passing (96.1%)"),
    ]
    for k, v in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        rk = p.add_run(f"{k}: ")
        rk.font.bold = True
        rk.font.size = Pt(11)
        rk.font.color.rgb = CLR_GREY
        rv = p.add_run(v)
        rv.font.size = Pt(11)
        rv.font.color.rgb = CLR_DARK

    doc.add_page_break()

    # ====================================================================
    # TABLE OF CONTENTS
    # ====================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Architecture Overview",
        "3. v3.4 Feature: Free LLM Tool Calling",
        "4. 12-Provider LLM Routing Matrix",
        "5. Bug Fixes & QC Remediation",
        "6. Nova Chatbot: 26 Tools",
        "7. API Integrations (30 Sources)",
        "8. Autonomous QC Engine (76 Tests)",
        "9. Deployment & Infrastructure",
        "10. Security & Safeguards",
        "11. Version History",
        "12. Recommendations",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = CLR_JOVEO_BLUE

    doc.add_page_break()

    # ====================================================================
    # 1. EXECUTIVE SUMMARY
    # ====================================================================
    add_heading_num(doc, "1", "Executive Summary")

    add_body(
        doc,
        "The Media Plan Generator v3.4 represents a major cost-optimization milestone. "
        "The primary change in this release is the migration of Nova chatbot's tool-calling "
        "pipeline from Claude Sonnet (paid, ~$0.003-0.005/query) to free LLM providers "
        "(Groq, Cerebras, SambaNova, xAI, Mistral, OpenRouter, NVIDIA NIM, Cloudflare). "
        "Claude is now used exclusively as a last-resort fallback, reducing per-query "
        "costs to near-zero for the majority of tool-calling interactions.",
    )

    add_body(
        doc,
        "This release builds on v3.3's deep QC audit (76 tests, 8-task routing matrix) "
        "and v3.2's 12-provider LLM router. The system now handles all three query types -- "
        "conversational, tool-calling, and verification -- through free providers first, "
        "with paid APIs (Claude Sonnet, Claude Opus, GPT-4o) available only when free "
        "providers fail or return insufficient quality.",
    )

    add_body(doc, "Key metrics for v3.4:")
    metrics = [
        ("LLM Providers", "12 total (9 free + 3 paid)"),
        ("Tool Definitions", "26 registered tools across recruitment marketing domain"),
        (
            "API Integrations",
            "30 live data sources (salary, demand, location, ad platforms)",
        ),
        ("QC Tests", "76 total, 73 passing (96.1%), 3 env-only failures"),
        ("Codebase", "25 Python modules, 61,219 lines of code"),
        (
            "Cost Reduction",
            "Tool-calling queries moved from ~$0.003-0.005/query to $0/query",
        ),
        ("Deployment", "Render.com Standard tier, auto-deploy from GitHub main branch"),
    ]
    for k, v in metrics:
        add_kv(doc, k, v)

    doc.add_page_break()

    # ====================================================================
    # 2. ARCHITECTURE OVERVIEW
    # ====================================================================
    add_heading_num(doc, "2", "Architecture Overview")

    add_body(
        doc,
        "The system is a standalone Python HTTP server (no Flask/Django) using only "
        "the standard library for HTTP handling. It provides REST API endpoints for "
        "media plan generation, AI chatbot (Nova), Slack bot integration, and comprehensive "
        "health/monitoring endpoints.",
    )

    add_heading_num(doc, "2.1", "Core Modules", level=2)

    modules = [
        (
            "app.py",
            "9,789",
            "HTTP server, all API endpoints, rate limiting, CORS, static files",
        ),
        (
            "nova.py",
            "5,213",
            "AI chatbot engine, 26 tools, Claude/free LLM integration",
        ),
        (
            "llm_router.py",
            "1,095",
            "12-provider LLM router, circuit breaker, task classification",
        ),
        (
            "api_enrichment.py",
            "10,502",
            "30 API integrations, concurrent fetching, caching",
        ),
        (
            "auto_qc.py",
            "3,064",
            "76-test autonomous QC engine, self-healing, self-upgrading",
        ),
        ("monitoring.py", "940", "Health checks, SLO tracking, metrics, audit trails"),
        ("ppt_generator.py", "4,845", "PowerPoint media plan generation"),
        (
            "research.py",
            "3,244",
            "Market research, knowledge base, 42 industry sources",
        ),
        ("data_orchestrator.py", "2,815", "Unified data fetching, tier-aware fallback"),
        (
            "budget_engine.py",
            "2,679",
            "Budget allocation, ROI projections, channel optimization",
        ),
        (
            "data_synthesizer.py",
            "2,532",
            "Multi-source data fusion, confidence scoring",
        ),
        ("nova_slack.py", "1,556", "Slack bot event handling, message formatting"),
        (
            "collar_intelligence.py",
            "1,620",
            "Blue/white collar classification and strategy",
        ),
        ("trend_engine.py", "1,360", "CPC/CPA trends, seasonal patterns, forecasting"),
        ("standardizer.py", "1,248", "Data normalization, schema validation"),
        ("generate_qc_audit_docx.py", "1,303", "QC audit report generation (DOCX)"),
        ("eval_framework.py", "1,227", "LLM evaluation framework"),
        ("data_contracts.py", "1,226", "Data validation contracts"),
        ("generate_product_docx.py", "927", "Product documentation generation"),
        ("regression_detector.py", "886", "Quality regression detection"),
        ("data_matrix_monitor.py", "743", "Data source health monitoring"),
        ("email_alerts.py", "701", "Email alerting via Resend"),
        ("grafana_logger.py", "535", "Grafana Loki log shipping"),
        ("supabase_cache.py", "915", "Supabase persistent caching"),
        ("shared_utils.py", "254", "Shared utility functions"),
    ]

    make_table(
        doc,
        ["Module", "Lines", "Purpose"],
        modules,
        col_widths=[2.0, 0.6, 4.1],
    )

    add_heading_num(doc, "2.2", "Request Flow", level=2)

    add_body(doc, "User Query Processing (4-tier routing):")
    add_bullet(
        doc,
        "Path A: Simple/conversational queries -> Free LLM providers (no tools, $0)",
    )
    add_bullet(doc, "Path B: Tool-use queries -> Free LLM providers WITH tools ($0)")
    add_bullet(
        doc, "Path C: Claude API -> Last-resort paid fallback (only if Path A/B fail)"
    )
    add_bullet(
        doc, "Path D: Rule-based fallback -> keyword matching + curated knowledge base"
    )

    add_body(doc, "Media Plan Generation Pipeline:")
    add_bullet(
        doc,
        "POST /api/generate -> async job creation -> data enrichment (30 APIs) -> "
        "LLM synthesis -> Excel/PPT generation -> download via /api/jobs/{id}",
    )

    doc.add_page_break()

    # ====================================================================
    # 3. v3.4 FEATURE: FREE LLM TOOL CALLING
    # ====================================================================
    add_heading_num(doc, "3", "v3.4 Feature: Free LLM Tool Calling")

    add_body(
        doc,
        "The headline feature of v3.4 is the migration of Nova's tool-calling pipeline "
        "from Claude Sonnet (Anthropic, paid) to free LLM providers. Previously, any query "
        "requiring data lookups (CPC benchmarks, salary data, location profiles, etc.) was "
        "forced to Claude regardless of cost. Now, 8 free OpenAI-compatible providers handle "
        "tool calling first, with Claude available only as a last resort.",
    )

    add_heading_num(doc, "3.1", "Technical Implementation", level=2)

    add_body(doc, "Changes to llm_router.py:")
    changes_router = [
        (
            "_convert_tools_anthropic_to_openai()",
            "New helper function that converts Anthropic "
            "tool format (name + input_schema) to OpenAI function-calling format (type: function, "
            "function: {name, parameters}). Strips Anthropic-specific keys like cache_control.",
        ),
        (
            "_build_openai_request()",
            "Extended to accept tools parameter. Handles tool result "
            "messages (role=tool), assistant messages with tool_calls, and adds converted tools "
            "to payload with tool_choice=auto.",
        ),
        (
            "_parse_openai_response()",
            "Extended to detect and return tool_calls from response. "
            "Preserves raw_message for conversation threading.",
        ),
        (
            "call_llm() routing",
            "Removed forced Claude routing for tools (was: if tools: "
            "force_provider = CLAUDE). Updated success detection to include tool_calls.",
        ),
    ]
    for name, desc in changes_router:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        rn = p.add_run(f"{name} -- ")
        rn.font.bold = True
        rn.font.size = Pt(10)
        rd = p.add_run(desc)
        rd.font.size = Pt(10)

    add_body(doc, "Changes to nova.py:")
    changes_nova = [
        (
            "_chat_with_free_llm_tools()",
            "New method (~150 lines). Multi-turn tool iteration "
            "loop with max 3 iterations. Provider lock-in ensures same provider handles all "
            "iterations. Includes paid-provider guard, JSON parse safety, quality gate.",
        ),
        (
            "_tool_handler_map() / _get_tool_handler_names()",
            "Centralized tool handler registry. "
            "Eliminates duplicated tool name validation. Used by both execute_tool() and "
            "free LLM tool path.",
        ),
        (
            "chat() routing",
            "Updated 4-tier routing: simple -> free+tools -> Claude -> rule-based. "
            "Free LLM with tools is now Path B, tried before Claude (Path C).",
        ),
        (
            "_FREE_TOOL_PROVIDERS",
            "Constant listing 8 free providers that support OpenAI-compatible "
            "tool calling: groq, cerebras, mistral, xai, sambanova, openrouter, nvidia_nim, cloudflare.",
        ),
    ]
    for name, desc in changes_nova:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        rn = p.add_run(f"{name} -- ")
        rn.font.bold = True
        rn.font.size = Pt(10)
        rd = p.add_run(desc)
        rd.font.size = Pt(10)

    add_heading_num(doc, "3.2", "Tool Format Conversion", level=2)

    add_body(
        doc,
        "Anthropic and OpenAI use different tool definition formats. The router automatically "
        "converts between them:",
    )

    make_table(
        doc,
        ["Aspect", "Anthropic Format", "OpenAI Format"],
        [
            (
                "Wrapper",
                "{ name, description, input_schema }",
                '{ type: "function", function: { name, description, parameters } }',
            ),
            ("Schema key", "input_schema", "parameters"),
            ("Extra keys", "cache_control (stripped)", "None"),
            ("Tool choice", "Not in payload", 'tool_choice: "auto"'),
            ("Response", "content[].type=tool_use", "message.tool_calls[].function"),
            ("Arguments", "content[].input (dict)", "function.arguments (JSON string)"),
        ],
        col_widths=[1.2, 2.5, 3.0],
    )

    add_heading_num(doc, "3.3", "Cost Impact", level=2)

    make_table(
        doc,
        ["Query Type", "Before v3.4", "After v3.4", "Savings"],
        [
            ("Simple/conversational", "$0 (free LLM)", "$0 (free LLM)", "No change"),
            ("Tool-calling (1 iteration)", "~$0.003 (Claude)", "$0 (free LLM)", "100%"),
            (
                "Tool-calling (3 iterations)",
                "~$0.009 (Claude)",
                "$0 (free LLM)",
                "100%",
            ),
            (
                "Complex tool chains (8 iter)",
                "~$0.024 (Claude)",
                "$0 free / $0.024 fallback",
                "~90%",
            ),
            ("Verification", "$0 (Gemini)", "$0 (Gemini)", "No change"),
        ],
        col_widths=[2.2, 1.6, 1.8, 1.1],
    )

    doc.add_page_break()

    # ====================================================================
    # 4. 12-PROVIDER LLM ROUTING MATRIX
    # ====================================================================
    add_heading_num(doc, "4", "12-Provider LLM Routing Matrix")

    add_body(
        doc,
        "The LLM router supports 12 providers organized into free and paid tiers. "
        "Each provider has an independent circuit breaker (5 failures -> 60s cooldown), "
        "per-minute rate tracking, and automatic fallback to the next provider in priority order.",
    )

    add_heading_num(doc, "4.1", "Provider Configuration", level=2)

    providers = [
        (
            "1",
            "Gemini 2.0 Flash",
            "gemini-2.0-flash",
            "Free",
            "15 RPM",
            "Structured, JSON, verification",
        ),
        (
            "2",
            "Groq",
            "llama-3.3-70b-versatile",
            "Free",
            "30 RPM",
            "Conversational, tool calling",
        ),
        (
            "3",
            "Cerebras",
            "llama-3.3-70b",
            "Free",
            "30 RPM",
            "Hot spare (same model, different infra)",
        ),
        (
            "4",
            "Mistral Small",
            "mistral-small-latest",
            "Free",
            "30 RPM",
            "JSON, multilingual",
        ),
        (
            "5",
            "OpenRouter",
            "meta-llama/llama-4-maverick",
            "Free",
            "20 RPM",
            "Free model gateway",
        ),
        (
            "6",
            "xAI Grok",
            "grok-3-mini-fast",
            "Free*",
            "30 RPM",
            "Strong reasoning, 128K context",
        ),
        (
            "7",
            "SambaNova",
            "Meta-Llama-3.1-405B",
            "Free",
            "10 RPM",
            "Largest open model (405B)",
        ),
        (
            "8",
            "NVIDIA NIM",
            "meta/llama-3.1-70b-instruct",
            "Free",
            "15 RPM",
            "NVIDIA-optimized inference",
        ),
        (
            "9",
            "Cloudflare",
            "@cf/meta/llama-3.3-70b",
            "Free",
            "50 RPM",
            "Edge-distributed, 10K neurons/day",
        ),
        ("10", "GPT-4o", "gpt-4o", "Paid", "60 RPM", "Structured + general reasoning"),
        (
            "11",
            "Claude Sonnet",
            "claude-sonnet-4-20250514",
            "Paid",
            "50 RPM",
            "Tool chains, last resort",
        ),
        (
            "12",
            "Claude Opus",
            "claude-opus-4-20250514",
            "Paid",
            "20 RPM",
            "Emergency only, highest quality",
        ),
    ]

    make_table(
        doc,
        ["#", "Provider", "Model", "Tier", "Rate Limit", "Primary Use"],
        providers,
        col_widths=[0.3, 1.2, 1.8, 0.5, 0.7, 2.2],
    )

    add_heading_num(doc, "4.2", "Task Classification (8 Types)", level=2)

    tasks = [
        (
            "STRUCTURED",
            "Benchmark lookups, CPC/CPA queries, JSON output",
            "Gemini > Groq > Cerebras",
        ),
        (
            "CONVERSATIONAL",
            "Strategy advice, general Q&A, advisory",
            "Groq > Cerebras > Gemini",
        ),
        (
            "COMPLEX",
            "What-if scenarios, role decomposition, multi-step",
            "Groq > Cerebras > SambaNova",
        ),
        (
            "CODE",
            "Formula generation, calculations, transforms",
            "Gemini > Groq > Cerebras",
        ),
        (
            "VERIFICATION",
            "Fact-checking, grounding validation",
            "Gemini > Groq > Cerebras",
        ),
        (
            "RESEARCH",
            "Market research, geopolitical analysis",
            "Groq > Cerebras > SambaNova",
        ),
        (
            "NARRATIVE",
            "Long-form text, executive summaries",
            "Groq > Cerebras > SambaNova",
        ),
        ("BATCH", "High-throughput bulk operations", "Gemini > Groq > Cerebras"),
    ]

    make_table(
        doc,
        ["Task Type", "Description", "Provider Priority (Free Tier)"],
        tasks,
        col_widths=[1.3, 2.8, 2.6],
    )

    add_heading_num(doc, "4.3", "Circuit Breaker Pattern", level=2)

    add_body(
        doc,
        "Each provider has an independent circuit breaker with 5-failure threshold and "
        "60-second cooldown. When a provider trips its circuit breaker, all requests are "
        "automatically routed to the next provider in priority order. Rate tracking is "
        "per-minute with configurable RPM limits per provider.",
    )

    add_body(
        doc,
        "Global timeout budget: 60 seconds maximum for the entire fallback loop. "
        "Individual per-provider timeouts are dynamically capped to the remaining budget. "
        "Minimum remaining budget to start a new attempt: 5 seconds.",
    )

    doc.add_page_break()

    # ====================================================================
    # 5. BUG FIXES & QC REMEDIATION
    # ====================================================================
    add_heading_num(doc, "5", "Bug Fixes & QC Remediation")

    add_body(
        doc,
        "The v3.4 release included a parallel code review (Opus-grade) that identified "
        "13 issues. 5 were fixed immediately (1 CRITICAL, 1 HIGH, 2 MEDIUM, 1 LOW). "
        "The remaining issues were informational or deferred.",
    )

    bugs = [
        (
            "CRITICAL",
            "call_llm() treated tool_calls as failures",
            "The routing loop success check (if result.get('text')) failed when a provider "
            "returned tool_calls (text is empty/null). This caused cascading through ALL providers. "
            "Fixed by adding tool_calls to the success detection: "
            "bool(result.get('text') or result.get('raw_content') or result.get('tool_calls')).",
            "llm_router.py",
        ),
        (
            "HIGH",
            "Free tool path could fall through to paid providers",
            "preferred_providers puts free first, but call_llm() appends the standard routing order "
            "which includes Claude/GPT-4o. A successful call to a paid provider in the free path "
            "would silently consume paid tokens. Fixed by adding paid-provider guard: "
            "if active_provider in {'gpt4o', 'claude', 'claude_opus'}: return None.",
            "nova.py",
        ),
        (
            "MEDIUM",
            "active_provider set even on failure (provider lock-in bug)",
            "First call could fail but still set active_provider, causing all retries to target "
            "a broken provider. Fixed by adding early bail on force_provider failure.",
            "nova.py",
        ),
        (
            "MEDIUM",
            "Hardcoded tool name validation set could drift",
            "Duplicated 26 tool names in _chat_with_free_llm_tools vs execute_tool. Any new tool "
            "added to one but not the other would break. Fixed by creating _tool_handler_map() + "
            "_get_tool_handler_names() shared methods.",
            "nova.py",
        ),
        (
            "LOW",
            "No per-response tool_call count limit",
            "Free models could hallucinate 10+ tool calls in a single response, wasting rate limit "
            "budget. Fixed by capping at 5 tool calls per response: tool_calls = tool_calls[:5].",
            "nova.py",
        ),
    ]

    for sev, title, desc, file in bugs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rt = p.add_run(f"{title}")
        rt.font.bold = True
        rt.font.size = Pt(10)
        add_severity_badge(p, sev)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.left_indent = Cm(0.63)
        rf = p2.add_run(f"File: {file}")
        rf.font.size = Pt(9)
        rf.font.color.rgb = CLR_GREY

        p3 = doc.add_paragraph(desc)
        p3.paragraph_format.space_after = Pt(8)
        p3.paragraph_format.left_indent = Cm(0.63)
        for run in p3.runs:
            run.font.size = Pt(9.5)

    doc.add_page_break()

    # ====================================================================
    # 6. NOVA CHATBOT: 26 TOOLS
    # ====================================================================
    add_heading_num(doc, "6", "Nova Chatbot: 26 Tools")

    add_body(
        doc,
        "Nova provides 26 tool definitions for Claude and free LLM providers. These tools "
        "enable real-time data lookups across recruitment marketing dimensions: supply data, "
        "salary benchmarks, market demand, location intelligence, ad platform metrics, and "
        "strategic analysis.",
    )

    tools = [
        (
            "1",
            "query_global_supply",
            "Joveo global supply data by role/location/channel",
        ),
        ("2", "query_channels", "Recruitment channel recommendations and comparisons"),
        ("3", "query_publishers", "Publisher/job board intelligence"),
        ("4", "query_knowledge_base", "42-source recruitment knowledge base"),
        ("5", "query_salary_data", "BLS/DataUSA salary benchmarks"),
        ("6", "query_market_demand", "JOLTS/BLS labor market demand"),
        ("7", "query_budget_projection", "Budget allocation and ROI projections"),
        ("8", "query_location_profile", "Location demographics, cost of living"),
        ("9", "query_ad_platform", "Ad platform recommendations"),
        ("10", "query_linkedin_guidewire", "LinkedIn audience intelligence"),
        ("11", "query_platform_deep", "Deep platform benchmarks (CPC/CPM/CTR)"),
        ("12", "query_recruitment_benchmarks", "Industry recruitment KPIs"),
        ("13", "query_employer_branding", "Employer brand strategy"),
        ("14", "query_regional_market", "Regional labor market analysis"),
        ("15", "query_supply_ecosystem", "Supply ecosystem mapping"),
        ("16", "query_workforce_trends", "Workforce and employment trends"),
        ("17", "query_white_papers", "Industry white papers and research"),
        (
            "18",
            "suggest_smart_defaults",
            "Auto-detect budget, channel split, CPA from partial info",
        ),
        (
            "19",
            "query_employer_brand",
            "Company-specific brand intelligence (30+ employers)",
        ),
        ("20", "query_ad_benchmarks", "CPC/CPM/CTR by platform and industry"),
        (
            "21",
            "query_hiring_insights",
            "Hiring difficulty, salary competitiveness, peak windows",
        ),
        (
            "22",
            "query_collar_strategy",
            "Blue vs white collar hiring strategy comparison",
        ),
        ("23", "query_market_trends", "CPC/CPA trends, seasonal patterns, forecasting"),
        (
            "24",
            "query_role_decomposition",
            "Seniority-level breakdowns and hiring splits",
        ),
        ("25", "simulate_what_if", "Budget/channel what-if scenario simulation"),
        (
            "26",
            "query_skills_gap",
            "Skills availability and hiring difficulty analysis",
        ),
    ]

    make_table(
        doc,
        ["#", "Tool Name", "Description"],
        tools,
        col_widths=[0.4, 2.2, 4.1],
    )

    doc.add_page_break()

    # ====================================================================
    # 7. API INTEGRATIONS (30 SOURCES)
    # ====================================================================
    add_heading_num(doc, "7", "API Integrations (30 Sources)")

    add_body(
        doc,
        "The api_enrichment.py module (10,502 lines) integrates with 30 external data sources. "
        "All API calls use stdlib urllib.request, have 8-second timeouts, are cached in-memory "
        "and on disk (24-hour TTL), fail gracefully, and run concurrently via ThreadPoolExecutor "
        "(max 15 workers).",
    )

    apis = [
        ("1", "BLS OES", "Salary data", "Free/Key"),
        ("2", "BLS QCEW", "Industry employment & wages", "Free"),
        ("3", "BLS JOLTS", "Job openings, hires, quits", "Free/Key"),
        ("4", "US Census ACS", "Location demographics", "Free"),
        ("5", "World Bank", "Global economic indicators", "Free"),
        ("6", "FRED", "US economic indicators", "Key"),
        ("7", "FRED Employment", "Hourly earnings, ECI", "Key"),
        ("8", "IMF DataMapper", "International GDP, inflation", "Free"),
        ("9", "DataUSA", "US occupation wages", "Free"),
        ("10", "O*NET", "Occupation skills, outlook", "Free/Creds"),
        ("11", "Eurostat LFS", "EU unemployment, wages", "Free"),
        ("12", "ILO ILOSTAT", "Global labor data", "Free"),
        ("13", "H-1B Wages", "Prevailing wages by SOC", "Curated"),
        ("14", "REST Countries", "Population, currency, languages", "Free"),
        ("15", "GeoNames", "Geographic data, coordinates", "Free"),
        ("16", "Teleport", "Quality of life, cost of living", "Free"),
        ("17", "Clearbit Logo", "Company logos", "Free"),
        ("18", "Clearbit Autocomplete", "Company metadata", "Free"),
        ("19", "Google Favicons", "Company favicons", "Free"),
        ("20", "Wikipedia REST", "Company descriptions", "Free"),
        ("21", "SEC EDGAR", "Public company data", "Free"),
        ("22", "Adzuna", "Job postings & salary", "Key"),
        ("23", "CareerOneStop", "DOL salary, outlook", "Key"),
        ("24", "Jooble", "International job market (69 countries)", "Key"),
        ("25", "Google Ads API", "Keyword volumes, CPC/CPM", "OAuth2/Bench"),
        ("26", "Meta Marketing", "Audience sizing, CPC/CPM", "Token/Bench"),
        ("27", "Microsoft/Bing Ads", "Search volumes, CPC", "OAuth2/Bench"),
        ("28", "TikTok Marketing", "Audience estimation, CPC", "Token/Bench"),
        ("29", "LinkedIn Marketing", "Professional audience, CPC", "Token/Bench"),
        ("30", "Currency Rates", "Exchange rates (live + fallback)", "Free"),
    ]

    make_table(
        doc,
        ["#", "API Source", "Data Type", "Auth"],
        apis,
        col_widths=[0.4, 1.8, 2.5, 1.0],
    )

    doc.add_page_break()

    # ====================================================================
    # 8. AUTONOMOUS QC ENGINE
    # ====================================================================
    add_heading_num(doc, "8", "Autonomous QC Engine (76 Tests)")

    add_body(
        doc,
        "The auto_qc.py module (3,064 lines) is a self-running, self-healing, self-upgrading "
        "quality assurance system. It runs the full test suite every 12 hours, attempts "
        "auto-resolution on failures, and generates new test cases weekly from user interaction "
        "analysis.",
    )

    add_heading_num(doc, "8.1", "Test Categories", level=2)

    test_cats = [
        (
            "Core Functionality",
            "API endpoints, media plan generation, data enrichment",
            "15",
        ),
        (
            "Nova Chatbot",
            "Tool execution, response quality, conversation handling",
            "12",
        ),
        (
            "LLM Router",
            "Provider routing, fallback, circuit breaker, task classification",
            "10",
        ),
        ("Data Integrity", "Schema validation, data contracts, normalization", "8"),
        ("Security", "Input sanitization, rate limiting, CORS, XSS prevention", "7"),
        ("Performance", "Latency, memory usage, concurrent request handling", "6"),
        ("Monitoring", "Health endpoints, metrics, SLO tracking", "5"),
        ("Integration", "Slack bot, email alerts, Grafana logging", "5"),
        ("Regression", "Historical regression detection, quality baselines", "4"),
        (
            "Dynamic (auto-generated)",
            "Weekly auto-generated from user interactions",
            "4",
        ),
    ]

    make_table(
        doc,
        ["Category", "Coverage", "Tests"],
        test_cats,
        col_widths=[1.8, 3.5, 0.6],
    )

    add_heading_num(doc, "8.2", "Current Results", level=2)

    add_kv(doc, "Total Tests", "76")
    add_kv(doc, "Passing", "73 (96.1%)")
    add_kv(doc, "Failing", "3 (env-only, expected on local)")
    add_kv(doc, "Pass Rate Target", "95% (met)")

    add_body(
        doc,
        "The 3 failing tests are environment-specific: they require Render deployment "
        "environment variables (Slack tokens, Resend keys) that are not available in "
        "local development. These pass in production.",
    )

    add_heading_num(doc, "8.3", "Self-Healing Capabilities", level=2)

    healing = [
        "Cache invalidation on stale data detection",
        "Module hot-reload on import failures",
        "Sentinel value reset on corrupted state",
        "Automatic retry with exponential backoff",
        "Audit log persistence recovery",
    ]
    for item in healing:
        add_bullet(doc, item)

    doc.add_page_break()

    # ====================================================================
    # 9. DEPLOYMENT & INFRASTRUCTURE
    # ====================================================================
    add_heading_num(doc, "9", "Deployment & Infrastructure")

    add_heading_num(doc, "9.1", "Render.com Configuration", level=2)

    add_kv(doc, "Plan", "Standard Tier (paid)")
    add_kv(doc, "Auto-Deploy", "Enabled (GitHub main branch)")
    add_kv(doc, "Region", "US (Oregon)")
    add_kv(doc, "Build Command", "pip install -r requirements.txt")
    add_kv(doc, "Start Command", "python app.py")

    add_heading_num(doc, "9.2", "Environment Variables (27)", level=2)

    env_vars = [
        ("ANTHROPIC_API_KEY", "Claude Sonnet/Opus (paid, last resort)"),
        ("GEMINI_API_KEY", "Google Gemini 2.0 Flash"),
        ("GROQ_API_KEY", "Groq Llama 3.3 70B"),
        ("CEREBRAS_API_KEY", "Cerebras Llama 3.3 70B"),
        ("MISTRAL_API_KEY", "Mistral Small"),
        ("OPENROUTER_API_KEY", "OpenRouter (Llama 4 Maverick)"),
        ("XAI_API_KEY", "xAI Grok"),
        ("SAMBANOVA_API_KEY", "SambaNova Llama 3.1 405B"),
        ("NVIDIA_NIM_API_KEY", "NVIDIA NIM Llama 3.1 70B"),
        ("CLOUDFLARE_AI_TOKEN", "Cloudflare Workers AI"),
        ("CLOUDFLARE_ACCOUNT_ID", "Cloudflare account identifier"),
        ("OPENAI_API_KEY", "GPT-4o (paid fallback)"),
        ("BLS_API_KEY", "Bureau of Labor Statistics v2"),
        ("FRED_API_KEY", "Federal Reserve Economic Data"),
        ("ADMIN_API_KEY", "Admin endpoint authentication"),
        ("RESEND_API_KEY", "Email alerting via Resend"),
        ("SLACK_BOT_TOKEN", "Slack bot OAuth token"),
        ("SLACK_CLIENT_ID", "Slack app client ID"),
        ("SLACK_CLIENT_SECRET", "Slack app client secret"),
        ("SLACK_REFRESH_TOKEN", "Slack token refresh"),
        ("SLACK_SIGNING_SECRET", "Slack request verification"),
        ("SUPABASE_URL", "Supabase project URL"),
        ("SUPABASE_ANON_KEY", "Supabase anonymous key"),
        ("GRAFANA_API_KEY", "Grafana Cloud API key"),
        ("GRAFANA_LOKI_URL", "Grafana Loki endpoint"),
        ("GRAFANA_USER_ID", "Grafana user identifier"),
        ("ALERT_EMAIL_TO", "Alert notification recipient"),
    ]

    make_table(
        doc,
        ["Variable", "Purpose"],
        env_vars,
        col_widths=[2.2, 4.5],
    )

    add_heading_num(doc, "9.3", "API Endpoints", level=2)

    endpoints = [
        ("POST /api/generate", "Async media plan generation (returns job_id)"),
        ("GET /api/jobs/{id}", "Poll job status and download results"),
        ("POST /api/chat", "Nova AI chatbot (conversational + tool calling)"),
        ("POST /api/slack/events", "Slack bot event webhook"),
        ("GET /api/health", "Liveness check"),
        ("GET /api/health/ready", "Readiness check (all dependencies)"),
        ("GET /api/health/data-matrix", "Data source health matrix"),
        ("GET /api/health/auto-qc", "QC test results (admin)"),
        ("GET /api/health/slos", "SLO compliance dashboard"),
        ("GET /api/health/eval", "LLM evaluation results"),
        ("GET /api/metrics", "Prometheus-style metrics (admin)"),
        ("GET /api/nova/metrics", "Nova chatbot metrics"),
        ("GET /api/admin/usage", "API usage analytics (admin)"),
        ("GET /api/admin/keys", "API key management (admin)"),
        ("POST /api/admin/nova", "Nova admin controls"),
        ("GET /api/docs/openapi.json", "OpenAPI 3.0 specification"),
    ]

    make_table(
        doc,
        ["Endpoint", "Description"],
        endpoints,
        col_widths=[2.5, 4.2],
    )

    doc.add_page_break()

    # ====================================================================
    # 10. SECURITY & SAFEGUARDS
    # ====================================================================
    add_heading_num(doc, "10", "Security & Safeguards")

    add_heading_num(doc, "10.1", "Rate Limiting", level=2)
    add_bullet(
        doc, "Global rate limit on /api/chat to prevent distributed API cost abuse"
    )
    add_bullet(doc, "Per-provider RPM limits with automatic throttling")
    add_bullet(doc, "Per-IP rate tracking with configurable thresholds")

    add_heading_num(doc, "10.2", "Input Validation", level=2)
    add_bullet(doc, "Message length truncation (MAX_MESSAGE_LENGTH)")
    add_bullet(doc, "Security filter blocks internal/technical/exploit questions")
    add_bullet(doc, "Tool name validation against registered handler map")
    add_bullet(
        doc, "JSON parse safety for free LLM tool arguments (try/except with retry)"
    )
    add_bullet(
        doc,
        "Per-response tool_call count cap (max 5) prevents hallucinated tool floods",
    )

    add_heading_num(doc, "10.3", "Paid API Protection", level=2)
    add_bullet(
        doc,
        "Paid-provider guard in free tool path prevents accidental Claude/GPT-4o usage",
    )
    add_bullet(
        doc,
        "Claude is hardcoded as last-resort only (Path C, after all free providers)",
    )
    add_bullet(
        doc,
        "Max 3 tool iterations for free providers (vs 8 for Claude) to respect rate limits",
    )
    add_bullet(
        doc,
        "Force-provider failure bail: returns None immediately instead of cascading",
    )

    add_heading_num(doc, "10.4", "Observability", level=2)
    add_bullet(doc, "Grafana Loki log shipping for centralized log aggregation")
    add_bullet(doc, "Supabase persistent caching for cross-deploy state")
    add_bullet(doc, "Email alerts via Resend on critical failures")
    add_bullet(doc, "Slack notifications for QC failures and system events")
    add_bullet(doc, "SLO monitoring with error budget tracking")
    add_bullet(doc, "Request tracing with unique IDs per thread")

    doc.add_page_break()

    # ====================================================================
    # 11. VERSION HISTORY
    # ====================================================================
    add_heading_num(doc, "11", "Version History")

    versions = [
        (
            "3.4.0",
            "2026-03-10",
            "Free LLM tool calling, paid-provider guard, 5 bug fixes",
        ),
        (
            "3.3.0",
            "2026-03-08",
            "Deep QC audit, 8-task routing matrix, 76 tests, critical fixes",
        ),
        (
            "3.2.0",
            "2026-03-06",
            "12-provider LLM router, Resend/Grafana/Supabase integrations",
        ),
        (
            "3.1.0",
            "2026-03-04",
            "Production hardening, intelligence upgrades, LLM router",
        ),
        (
            "3.0.0",
            "2026-02-28",
            "AI Intelligence Engine, trend engine, collar intelligence",
        ),
    ]

    make_table(
        doc,
        ["Version", "Date", "Key Changes"],
        versions,
        col_widths=[0.8, 1.2, 4.7],
    )

    add_heading_num(doc, "11.1", "Git Commit Log (Recent)", level=2)

    commits = [
        ("0216724", "chore: bump version strings to 3.4.0 across all health endpoints"),
        (
            "9ee1c1c",
            "feat: v3.4 -- move Nova tool calling from Claude (paid) to free LLM providers",
        ),
        (
            "bea4e71",
            "feat: v3.3 -- deep QC audit, 8-task LLM routing matrix, 76 tests, critical fixes",
        ),
        ("85dc768", "feat: expand LLM router to 12 providers (9 free + 3 paid)"),
        (
            "c916bfc",
            "feat: v3.2 -- 9-provider LLM router, Resend/Grafana/Supabase integrations",
        ),
        ("5fd31b3", "fix: Nova chatbot accuracy, verbosity, and hallucination issues"),
        ("8b61f88", "docs: add v3.1 system state document (MD + DOCX)"),
        (
            "0afc021",
            "feat: v3.1 production hardening + intelligence upgrades + LLM router",
        ),
        (
            "ec5568b",
            "chore: bump to v3.0.0, delete joveo_iq.py, fix DataUSA latency waste",
        ),
        (
            "16afc76",
            "feat: v3 AI Intelligence Engine -- trend engine, collar intelligence",
        ),
    ]

    make_table(
        doc,
        ["Commit", "Message"],
        commits,
        col_widths=[0.9, 5.8],
    )

    doc.add_page_break()

    # ====================================================================
    # 12. RECOMMENDATIONS
    # ====================================================================
    add_heading_num(doc, "12", "Recommendations")

    add_heading_num(doc, "12.1", "Short-Term (Next Sprint)", level=2)
    recs_short = [
        "Monitor free LLM tool-calling success rates in production logs for 1-2 weeks before "
        "considering any adjustments to the provider priority order.",
        "Add structured logging for tool-calling path selection (free vs Claude fallback) "
        "to measure actual cost savings.",
        "Consider increasing max tool iterations from 3 to 4 for free providers if Groq/Cerebras "
        "prove reliable at handling multi-step chains.",
        "Add response quality scoring (automated) to compare free LLM tool responses vs Claude "
        "for the same queries.",
    ]
    for rec in recs_short:
        add_bullet(doc, rec)

    add_heading_num(doc, "12.2", "Medium-Term (1-3 Months)", level=2)
    recs_med = [
        "Implement streaming responses for the chatbot to improve perceived latency.",
        "Add Gemini to the tool-calling path (currently only in the no-tool path) -- Gemini "
        "supports function calling but uses a different format than OpenAI.",
        "Build a cost dashboard showing real-time spend across all 12 providers with "
        "trend visualization.",
        "Expand the dynamic test generation to use LLM analysis of production error logs.",
        "Add A/B testing framework to compare free vs paid provider response quality.",
    ]
    for rec in recs_med:
        add_bullet(doc, rec)

    add_heading_num(doc, "12.3", "Long-Term (3-6 Months)", level=2)
    recs_long = [
        "Evaluate fine-tuning a smaller model (Llama 3.3 8B) on Nova's tool-calling patterns "
        "for even faster inference with zero cost.",
        "Add multi-language support for the chatbot (leverage Mistral's multilingual strengths).",
        "Build a self-serve analytics dashboard for clients to explore their media plan data.",
        "Consider migrating from standalone HTTP server to a lightweight framework (FastAPI) "
        "for better async support and automatic OpenAPI generation.",
    ]
    for rec in recs_long:
        add_bullet(doc, rec)

    # ====================================================================
    # FOOTER
    # ====================================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("--- End of Report ---")
    run.font.color.rgb = CLR_GREY
    run.font.size = Pt(10)
    run.font.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Generated {now} | Media Plan Generator v3.4.0 | Joveo")
    run2.font.color.rgb = CLR_GREY
    run2.font.size = Pt(9)

    # ====================================================================
    # SAVE
    # ====================================================================
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Report saved to: {OUTPUT_FILE}")
    print(f"File size: {os.path.getsize(OUTPUT_FILE):,} bytes")
    return OUTPUT_FILE


if __name__ == "__main__":
    build_report()
