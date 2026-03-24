"""
generate_product_docx.py -- Generate the Nova AI Suite Product Overview DOCX.

Creates a professional 10-page product overview document using python-docx.
Run directly: python3 generate_product_docx.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
DARK_BLUE = RGBColor(0x00, 0x3D, 0x6B)  # Primary heading color
JOVEO_BLUE = RGBColor(0x00, 0x66, 0xCC)  # Accent / links
MEDIUM_BLUE = RGBColor(0x1A, 0x5C, 0x8E)  # Subheadings
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)  # Body text
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)  # Secondary text
TABLE_HEADER_BG = "003D6B"  # Hex for table header fill
TABLE_ALT_ROW_BG = "EAF0F6"  # Hex for alternating rows
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Joveo_Media_Plan_Generator_Product_Overview.docx",
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professional styled table with header row and alternating shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(hdr_cells[i], TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(row_cells[c_idx], TABLE_ALT_ROW_BG)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = width

    return table


def add_heading1(doc, text):
    """Add a Heading 1 styled paragraph."""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.name = "Calibri"
    return h


def add_heading2(doc, text):
    """Add a Heading 2 styled paragraph."""
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = MEDIUM_BLUE
        run.font.name = "Calibri"
    return h


def add_body(doc, text, bold=False, italic=False):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_GRAY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_prefix=""):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run_b.font.name = "Calibri"
        run_b.font.color.rgb = DARK_GRAY
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_GRAY
    return p


def add_page_break(doc):
    """Insert an explicit page break."""
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------


def build_document():
    doc = Document()

    # --- Global style defaults ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY

    # Adjust default section margins
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ===================================================================
    # PAGE 1: TITLE PAGE
    # ===================================================================

    # Add vertical spacing
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Nova AI Suite")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("AI-Powered Recruitment Advertising Intelligence Platform")
    run.font.size = Pt(16)
    run.font.color.rgb = MEDIUM_BLUE
    run.font.name = "Calibri"

    # Separator line
    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep_p.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(10)

    # Doc type
    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = type_p.add_run("Product Overview & Technical Architecture")
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_GRAY
    run.font.name = "Calibri"

    # Version
    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_p.add_run("v3.2  |  March 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = LIGHT_GRAY
    run.font.name = "Calibri"

    # Spacer
    doc.add_paragraph()
    doc.add_paragraph()

    # Company tagline
    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag_p.add_run(
        "Nova AI Suite -- Precision Programmatic for Talent Acquisition"
    )
    run.font.size = Pt(13)
    run.font.color.rgb = JOVEO_BLUE
    run.bold = True
    run.font.name = "Calibri"

    add_page_break(doc)

    # ===================================================================
    # PAGE 2: EXECUTIVE SUMMARY
    # ===================================================================

    add_heading1(doc, "Executive Summary")

    add_body(
        doc,
        "The Nova AI Suite Media Plan Generator is an AI-powered platform that transforms "
        "recruitment advertising planning from a weeks-long manual process into minutes "
        "of automated, data-driven analysis. It ingests client requirements -- industry, "
        "roles, locations, budget, and competitors -- and produces comprehensive media "
        "plans backed by real-time market intelligence from over 30 public and proprietary APIs.",
    )

    add_heading2(doc, "Four Delivery Channels")

    add_bullet(
        doc,
        " A browser-based interface that generates downloadable Excel "
        "spreadsheets and PowerPoint presentations as a bundled media plan.",
        bold_prefix="Web UI:",
    )
    add_bullet(
        doc,
        " An AI chatbot (Nova) accessible through the web UI, providing "
        "conversational access to recruitment intelligence, benchmark data, "
        "and strategic recommendations.",
        bold_prefix="Nova AI Chatbot:",
    )
    add_bullet(
        doc,
        " A workplace-native integration where team members can query "
        "Nova directly in Slack via mentions or direct messages, with full "
        "access to the same data intelligence layer.",
        bold_prefix="Slack Bot:",
    )
    add_bullet(
        doc,
        " Programmatic endpoints for headless media plan generation, "
        "health monitoring, and administrative operations.",
        bold_prefix="API:",
    )

    add_heading2(doc, "Key Platform Statistics")

    stats_data = [
        [
            "30+",
            "Real-time API integrations (BLS, Census, World Bank, FRED, O*NET, and more)",
        ],
        ["12", "LLM providers in the unified router (9 free-tier, 3 paid-tier)"],
        ["66+", "Automated QC tests running autonomously every 12 hours"],
        ["26", "Supported currencies with proper symbol formatting"],
        ["22", "Industry verticals with 4-year trend data across 6 ad platforms"],
        [
            "4",
            "Collar types classified (blue, white, grey, pink) for strategy differentiation",
        ],
        ["0", "External pip dependencies for the core server (Python stdlib only)"],
    ]
    add_styled_table(
        doc,
        ["Metric", "Description"],
        stats_data,
        col_widths=[Inches(0.8), Inches(5.2)],
    )

    add_page_break(doc)

    # ===================================================================
    # PAGE 3: PLATFORM CAPABILITIES
    # ===================================================================

    add_heading1(doc, "Platform Capabilities")

    add_heading2(doc, "Media Plan Generation (Excel + PowerPoint)")
    add_body(
        doc,
        "The core output pipeline accepts structured input (industry, roles, locations, "
        "budget, competitors) and produces a professional Excel workbook with per-channel "
        "budget allocations, projected clicks, applications, and hires, alongside a "
        "branded PowerPoint deck summarizing the strategic rationale. The generation "
        "pipeline runs API enrichment, data synthesis, budget optimization, and document "
        "formatting in a single coordinated flow.",
    )

    add_heading2(doc, "Nova AI Chatbot")
    add_body(
        doc,
        "Nova is an AI-powered recruitment marketing intelligence chatbot embedded "
        "in the web UI. It provides conversational access to Nova AI Suite's proprietary supply "
        "data (10,238+ publishers across 70+ countries), 30+ live API enrichment sources, "
        "a 42-source recruitment industry knowledge base, the data synthesis engine with "
        "confidence scoring, and the budget allocation engine with dollar projections. "
        "Nova operates in two modes: rule-based keyword routing for fast deterministic "
        "answers, and LLM-powered reasoning via Claude Haiku 4.5 (simple queries) or "
        "Claude Sonnet 4.6 (complex strategy queries).",
    )

    add_heading2(doc, "Slack Bot Integration")
    add_body(
        doc,
        "The Nova Slack Bot brings the full intelligence layer into the workplace. Team "
        "members can mention @Nova or send direct messages in Slack to access all 21 Nova "
        "tools, including orchestrator-powered data enrichment. The bot searches Slack "
        "history for previously answered questions, maintains an unanswered question queue "
        "for human review, learns from human-provided answers over time, and sends weekly "
        "digest summaries of outstanding questions. Built entirely with Python stdlib "
        "(urllib.request, no slack_sdk dependency).",
    )

    add_heading2(doc, "Budget Optimization Engine")
    add_body(
        doc,
        "The budget engine converts percentage allocations to concrete dollar amounts, "
        "producing per-channel spend recommendations with projected outcomes (clicks, "
        "applications, hires). It integrates dynamic CPC benchmarks from the trend engine, "
        "collar-weighted allocation via the collar intelligence module, and structured "
        "confidence on every channel allocation. Role-tier multipliers adjust CPA "
        "expectations across six seniority levels from Gig (0.5x) to Executive (3.5x).",
    )

    add_heading2(doc, "Multi-Currency & Global Support")
    add_body(
        doc,
        "The platform supports 26 currencies (USD, EUR, GBP, INR, JPY, CNY, AUD, CAD, "
        "SGD, HKD, NZD, CHF, SEK, NOK, DKK, BRL, ZAR, MXN, KRW, THB, MYR, PHP, IDR, "
        "AED, SAR, and more) with proper symbol formatting. Location intelligence spans "
        "40+ countries and 100+ metro areas, with automatic currency detection based on "
        "the primary hiring location.",
    )

    add_heading2(doc, "Collar Intelligence")
    add_body(
        doc,
        "First-class collar-type classification differentiates strategy at every level: "
        "data queries, budget allocation, channel selection, output formatting, and chat "
        "responses. The engine classifies roles into four collar types (blue, white, grey, "
        "pink) using SOC major groups, O*NET Job Zones, standardizer role tiers, and "
        "keyword pattern matching. Each collar type routes to different channel mixes, "
        "budget weightings, and messaging strategies.",
    )

    add_page_break(doc)

    # ===================================================================
    # PAGE 4: THE AI ENGINE -- UNIFIED LLM ROUTER
    # ===================================================================

    add_heading1(doc, "The AI Engine: Unified LLM Router")

    add_body(
        doc,
        "At the core of Nova's intelligence is a unified LLM router that manages twelve "
        "independent AI providers through a cost-optimized, fault-tolerant cascade. The "
        "router follows a free-first philosophy: it exhausts all nine free-tier providers "
        "before falling back to three paid models, minimizing operational cost while maintaining "
        "response quality. The entire routing layer is stdlib-only and thread-safe.",
    )

    add_heading2(doc, "Provider Cascade (Priority Order)")

    provider_rows = [
        [
            "1",
            "Gemini 2.0 Flash",
            "Google",
            "Free",
            "Structured JSON, code, verification",
        ],
        [
            "2",
            "Groq Llama 3.3 70B",
            "Groq",
            "Free",
            "Conversational, complex reasoning",
        ],
        [
            "3",
            "Cerebras Llama 3.3 70B",
            "Cerebras",
            "Free",
            "Hot spare (same model, independent infra)",
        ],
        [
            "4",
            "Mistral Small",
            "Mistral AI",
            "Free",
            "JSON output, multilingual support",
        ],
        [
            "5",
            "OpenRouter (Llama 4 Maverick)",
            "OpenRouter",
            "Free",
            "Strong general purpose via free models",
        ],
        ["6", "xAI Grok", "xAI", "Free*", "Strong reasoning ($25 signup credits)"],
        [
            "7",
            "SambaNova (Llama 3.1 405B)",
            "SambaNova",
            "Free",
            "Largest open model free, fastest inference (RDU)",
        ],
        [
            "8",
            "NVIDIA NIM (Llama 3.1 70B)",
            "NVIDIA",
            "Free",
            "NVIDIA-optimized, exclusive Nemotron models",
        ],
        [
            "9",
            "Cloudflare Workers AI (Llama 3.3 70B)",
            "Cloudflare",
            "Free",
            "Edge-distributed, 10K neurons/day",
        ],
        [
            "10",
            "GPT-4o",
            "OpenAI",
            "Paid",
            "Structured JSON, general reasoning, calculations",
        ],
        [
            "11",
            "Claude Sonnet 4",
            "Anthropic",
            "Paid",
            "High quality, strong tool-use chains",
        ],
        [
            "12",
            "Claude Opus 4",
            "Anthropic",
            "Paid",
            "Last resort, highest quality, most expensive",
        ],
    ]
    add_styled_table(
        doc,
        ["#", "Provider", "Vendor", "Tier", "Strength"],
        provider_rows,
        col_widths=[
            Inches(0.35),
            Inches(2.0),
            Inches(0.85),
            Inches(0.45),
            Inches(2.35),
        ],
    )

    add_heading2(doc, "Task-Based Routing")

    add_body(
        doc,
        "Every incoming request is classified into one of four task types, and the "
        "provider priority order is adjusted accordingly within each tier:",
    )

    routing_rows = [
        [
            "STRUCTURED",
            "Benchmark lookups, CPC/CPA queries, JSON output",
            "Gemini > Mistral > Groq > ... > GPT-4o",
        ],
        [
            "CONVERSATIONAL",
            "Strategy explanations, general Q&A, advisory",
            "Groq > Cerebras > Gemini > ... > GPT-4o",
        ],
        [
            "COMPLEX",
            "What-if scenarios, role decomposition, multi-step analysis",
            "Groq > Cerebras > Gemini > ... > Claude Sonnet",
        ],
        [
            "CODE",
            "Formula generation, calculations, data transforms",
            "Gemini > Mistral > Groq > ... > GPT-4o",
        ],
    ]
    add_styled_table(
        doc,
        ["Task Type", "Description", "Preferred Providers"],
        routing_rows,
        col_widths=[Inches(1.2), Inches(2.8), Inches(2.0)],
    )

    add_heading2(doc, "Fault Tolerance")

    add_bullet(
        doc,
        " Each provider has an independent circuit breaker. After 5 consecutive "
        "failures, the provider enters a 60-second cooldown before retrying.",
        bold_prefix="Circuit Breakers:",
    )
    add_bullet(
        doc,
        " Per-minute and per-day request tracking per provider prevents "
        "exceeding API quotas (e.g., Gemini: 15 RPM / 1,500 RPD; Groq: 30 RPM / 14,400 RPD).",
        bold_prefix="Rate Limiting:",
    )
    add_bullet(
        doc,
        " The entire call_llm() fallback loop is capped at 60 seconds "
        "wall-clock time. Individual provider timeouts are dynamically adjusted to fit "
        "within the remaining budget. No attempt starts with less than 5 seconds remaining.",
        bold_prefix="Global Timeout Budget:",
    )

    add_page_break(doc)

    # ===================================================================
    # PAGE 5: DATA INTELLIGENCE LAYER
    # ===================================================================

    add_heading1(doc, "Data Intelligence Layer")

    add_body(
        doc,
        "The platform's data intelligence layer is built on a unified data orchestrator "
        "that cascades through all available data sources in order of cost and speed. "
        "It never crashes -- all errors are caught and the caller always receives a "
        "usable result dictionary.",
    )

    add_heading2(doc, "Unified Data Orchestrator")
    add_body(
        doc,
        "The orchestrator (data_orchestrator.py) is the single entry point for enriched "
        "data queries. It cascades through six tiers:",
    )

    add_bullet(
        doc,
        " research.py embedded data (free, instant, 40+ countries, 100+ metros)",
        bold_prefix="Tier 1:",
    )
    add_bullet(
        doc,
        " trend_engine.py benchmarks (free, instant, 4-year CPC/CPA trends)",
        bold_prefix="Tier 2:",
    )
    add_bullet(
        doc,
        " collar_intelligence.py (free, instant, collar classification)",
        bold_prefix="Tier 3:",
    )
    add_bullet(
        doc,
        " Selective live API calls (individual APIs, cached 24 hours)",
        bold_prefix="Tier 4:",
    )
    add_bullet(
        doc,
        " data_synthesizer.py fusion (cross-validates multi-source data)",
        bold_prefix="Tier 5:",
    )
    add_bullet(
        doc, " Static KB fallback (JSON files, always available)", bold_prefix="Tier 6:"
    )

    add_heading2(doc, "30+ Real-Time API Integrations")

    api_rows = [
        ["BLS OES", "Bureau of Labor Statistics", "Salary data by occupation"],
        [
            "BLS QCEW",
            "Bureau of Labor Statistics",
            "Industry employment and wage statistics",
        ],
        [
            "BLS JOLTS",
            "Bureau of Labor Statistics",
            "Job openings, hires, quits by industry",
        ],
        [
            "US Census ACS",
            "U.S. Census Bureau",
            "Location demographics, population, income",
        ],
        ["O*NET", "Dept. of Labor", "Occupation skills, knowledge, outlook"],
        ["CareerOneStop", "Dept. of Labor", "Salary, outlook, certifications"],
        ["FRED", "Federal Reserve", "Economic indicators, avg hourly earnings, ECI"],
        [
            "World Bank",
            "World Bank Group",
            "Global economic indicators (190+ countries)",
        ],
        [
            "IMF DataMapper",
            "International Monetary Fund",
            "GDP, inflation, unemployment",
        ],
        ["Eurostat LFS", "European Commission", "EU unemployment, wages, employment"],
        [
            "ILO ILOSTAT",
            "International Labour Org.",
            "Global labour participation rates",
        ],
        ["SEC EDGAR", "Securities & Exchange Comm.", "Public company filings and data"],
        ["Google Ads API", "Google", "Keyword search volumes, CPC/CPM benchmarks"],
        ["Meta Marketing API", "Meta", "Facebook/Instagram audience sizing, CPC/CPM"],
        ["LinkedIn Marketing API", "LinkedIn", "Professional audience sizing, CPC"],
        ["Bing Ads API", "Microsoft", "Search volumes, CPC estimates"],
        ["TikTok Marketing API", "TikTok / ByteDance", "Audience estimation, CPC/CPM"],
        ["Adzuna", "Adzuna", "Job postings and salary data"],
        ["Jooble", "Jooble", "International job market data (69 countries)"],
        ["REST Countries", "Open Source", "Country population, currency, languages"],
        ["GeoNames", "GeoNames.org", "Geographic data, coordinates, timezone"],
        ["Teleport", "Teleport", "Quality of life scores, cost of living"],
        ["DataUSA", "Datawheel / Deloitte", "US occupation wages, state demographics"],
        ["Google Trends", "Google", "Search interest data for job market signals"],
        ["Wikipedia REST", "Wikimedia Foundation", "Company descriptions"],
        ["Clearbit", "Clearbit / HubSpot", "Company metadata and domain lookup"],
        ["Currency Rates", "Various", "Live exchange rates (26 currencies)"],
        ["H-1B Wage Data", "DOL / Curated", "Prevailing wages by SOC code"],
    ]
    add_styled_table(
        doc,
        ["API", "Provider", "Data Provided"],
        api_rows,
        col_widths=[Inches(1.3), Inches(1.7), Inches(3.0)],
    )

    add_heading2(doc, "3-Tier Caching Architecture")

    add_bullet(
        doc,
        " In-memory dictionary -- fastest, lost on restart.",
        bold_prefix="L1 (Memory):",
    )
    add_bullet(
        doc,
        " Disk JSON files -- survives restart, limited by available disk.",
        bold_prefix="L2 (Disk):",
    )
    add_bullet(
        doc,
        " Supabase Postgres -- survives redeployments, shared across instances. "
        "Uses only urllib.request with 3-second timeout per call, retry on 5xx errors, "
        "and automatic TTL-based cleanup.",
        bold_prefix="L3 (Supabase):",
    )

    add_heading2(doc, "Canonical Taxonomy Standardizer")
    add_body(
        doc,
        "The standardizer module (standardizer.py) provides unified naming conventions "
        "across all subsystems. It normalizes industries (17 canonical entries), roles, "
        "locations (country maps, US state maps, region maps), platforms/channels, and "
        "metrics. Every public function handles None/empty input gracefully and performs "
        "case-insensitive matching, ensuring consistent data flow between the frontend, "
        "API enrichment, data synthesizer, knowledge base files, Nova, and channel databases.",
    )

    add_page_break(doc)

    # ===================================================================
    # PAGE 6: BUDGET ENGINE & COLLAR INTELLIGENCE
    # ===================================================================

    add_heading1(doc, "Budget Engine & Collar Intelligence")

    add_heading2(doc, "Budget Optimization Algorithm")
    add_body(
        doc,
        "The budget engine (budget_engine.py) converts high-level percentage allocations "
        "into concrete dollar amounts with projected outcomes. It operates on a hierarchy "
        "of data sources: dynamic CPC benchmarks from the trend engine take priority, "
        "followed by collar-weighted adjustments from collar intelligence, then static "
        "base benchmarks as a final fallback.",
    )

    add_heading2(doc, "Role-Tier CPA Multipliers")
    add_body(
        doc,
        "Cost-per-acquisition expectations are adjusted by role seniority using predefined "
        "multipliers. These multipliers scale the base CPA to reflect the realistic cost "
        "of filling roles at different levels:",
    )

    tier_rows = [
        ["Executive / Leadership", "3.5x", "C-suite, VP, Director roles"],
        ["Clinical / Licensed", "2.2x", "Nurses, physicians, licensed practitioners"],
        ["Professional / White-Collar", "1.8x", "Engineers, analysts, managers"],
        [
            "Skilled Trades / Technical",
            "1.0x",
            "Electricians, mechanics, technicians (baseline)",
        ],
        ["Education / Academic", "1.0x", "Teachers, professors, academic staff"],
        ["Hourly / Entry-Level", "0.7x", "Retail, food service, warehouse workers"],
        ["Gig / Independent Contractor", "0.5x", "Freelancers, on-demand workers"],
    ]
    add_styled_table(
        doc,
        ["Role Tier", "Multiplier", "Example Roles"],
        tier_rows,
        col_widths=[Inches(1.8), Inches(0.8), Inches(3.4)],
    )

    add_heading2(doc, "Collar-Aware Channel Optimization")
    add_body(
        doc,
        "The collar intelligence engine (collar_intelligence.py) classifies every role "
        "into one of four collar types using multiple signals: SOC major group codes from "
        "BLS (e.g., groups 47-53 map to blue collar, 11-29 to white collar), O*NET Job "
        "Zones (1-2 = blue/entry, 3 = grey/skilled, 4-5 = white/professional), standardizer "
        "role tiers, and keyword pattern matching for unclassified roles.",
    )

    collar_rows = [
        [
            "Blue Collar",
            "Manual labor, trades, hourly, shift-based",
            "Mobile-first, job boards, programmatic",
        ],
        [
            "White Collar",
            "Professional, office-based, salaried",
            "LinkedIn, search, professional networks",
        ],
        [
            "Grey Collar",
            "Licensed/clinical, shift-based (nurses, techs)",
            "Niche healthcare boards, hybrid channels",
        ],
        [
            "Pink Collar",
            "Administrative, care, service roles",
            "General boards, social, community channels",
        ],
    ]
    add_styled_table(
        doc,
        ["Collar Type", "Description", "Preferred Channels"],
        collar_rows,
        col_widths=[Inches(1.0), Inches(2.5), Inches(2.5)],
    )

    add_heading2(doc, "Trend Engine: Dynamic Benchmarks")
    add_body(
        doc,
        "The trend engine (trend_engine.py) is the single source of truth for recruitment "
        "advertising benchmark data. It provides 4-year historical CPC/CPM/CTR/CPA trends "
        "across 6 ad platforms (Google Search, Meta Facebook, Meta Instagram, LinkedIn, "
        "Indeed, Programmatic) and 22 industry verticals. Features include seasonal monthly "
        "multipliers differentiated by collar type, regional CPC adjustment factors for "
        "100+ US metros and 40+ countries, collar-type CPC differentials per platform, and "
        "structured uncertainty on every returned value (confidence interval, trend direction). "
        "Data is curated from Appcast, WordStream/LOCALiQ, SHRM, LinkedIn Talent Solutions, "
        "Recruitics/PandoLogic, and iCIMS benchmark reports spanning 2022-2026.",
    )

    add_page_break(doc)

    # ===================================================================
    # PAGE 7: QUALITY ASSURANCE & SELF-HEALING
    # ===================================================================

    add_heading1(doc, "Quality Assurance & Self-Healing")

    add_heading2(doc, "AutoQC Engine")
    add_body(
        doc,
        "The autonomous QC engine (auto_qc.py) is a self-running, self-healing, "
        "self-upgrading quality assurance system. It executes a comprehensive test suite "
        "against live endpoints on a twice-daily schedule (every 12 hours), with a 45-second "
        "hard ceiling per individual test. Results are stored in memory and persisted to disk, "
        "with failures triggering self-healing actions and optional Slack alerts.",
    )

    add_bullet(
        doc,
        " Static tests ported from the original test_nova_chat.sh script, "
        "covering endpoint health, response format validation, data integrity, "
        "and integration correctness.",
        bold_prefix="Static Tests:",
    )
    add_bullet(
        doc,
        " Generated weekly from request_log.json and Nova interaction metrics. "
        "The engine analyzes real user interactions to create new test cases that "
        "cover emerging usage patterns.",
        bold_prefix="Dynamic Tests:",
    )
    add_bullet(
        doc,
        " On failure, the engine attempts automatic resolution: cache clearing, "
        "module reload via importlib, data orchestrator sentinel resets, and stale "
        "cache eviction.",
        bold_prefix="Self-Healing:",
    )
    add_bullet(
        doc,
        " Every 7 days, the engine analyzes user interaction patterns to "
        "generate new test cases, keeping the test suite aligned with actual "
        "production usage.",
        bold_prefix="Weekly Self-Upgrade:",
    )

    add_heading2(doc, "Data Matrix Monitor")
    add_body(
        doc,
        "The data matrix monitor (data_matrix_monitor.py) tracks whether all four products "
        "(Excel/PPT, Nova Chat, Slack Bot, PPT Generator) correctly access all nine data "
        "layers. It runs probes every 12 hours in a background daemon thread and maps an "
        "expected-state matrix (YES / NO / PARTIAL / VIA_ORCHESTRATOR) against actual "
        "runtime conditions. Self-healing actions include re-importing failed modules, "
        "resetting orchestrator lazy-load sentinels, and evicting stale cache entries.",
    )

    add_heading2(doc, "Evaluation Framework")
    add_body(
        doc,
        "The evaluation framework (eval_framework.py) scores the quality of budget "
        "recommendations, collar classification consistency, geographic coherence, and "
        "CPA reasonableness across the core modules (budget_engine, collar_intelligence, "
        "trend_engine). Every test case execution is wrapped in try/except so a single "
        "failure cannot abort the suite.",
    )

    add_heading2(doc, "Regression Detector")
    add_body(
        doc,
        "The regression detector (regression_detector.py) runs a fixed set of reference "
        "scenarios through the budget allocation engine and compares outputs to a persisted "
        "baseline snapshot. Alert thresholds are configurable: CPA drift greater than 10%, "
        "channel allocation drift greater than 15%, and hire projection drift greater than "
        "20% all trigger alerts. The baseline file is persisted at "
        "data/persistent/regression_baseline.json.",
    )

    add_heading2(doc, "Data Contracts")
    add_body(
        doc,
        "The data contracts module (data_contracts.py) validates 12 JSON knowledge-base "
        "files and API enrichment output against defined schemas, ensuring that downstream "
        "consumers (PPT generator, chatbot, orchestrator) never silently receive malformed "
        "data. Schema validation covers structure, required keys, numeric ranges, and type "
        "correctness.",
    )

    add_page_break(doc)

    # ===================================================================
    # PAGE 8: OBSERVABILITY & MONITORING
    # ===================================================================

    add_heading1(doc, "Observability & Monitoring")

    add_heading2(doc, "Structured Logging & Request Tracing")
    add_body(
        doc,
        "The monitoring module (monitoring.py) provides structured JSON logging with "
        "request-scoped tracing. Each incoming request receives a unique 12-character hex "
        "request ID, and all log entries within that request's scope carry the ID for "
        "end-to-end correlation. Elapsed time tracking is built into the request context, "
        "enabling per-request latency measurement.",
    )

    add_heading2(doc, "Grafana Cloud Loki Integration")
    add_body(
        doc,
        "The grafana_logger.py module ships structured JSON logs to Grafana Cloud Loki "
        "for centralized observability. The handler attaches to the existing Python logging "
        "system with a default threshold of WARNING, so only warning, error, and critical "
        "records are shipped externally while debug and info logs stay local. Log entries "
        "include labels for application name, deployment environment, and severity level. "
        "The module is entirely stdlib-based and gracefully disables itself when environment "
        "variables are not configured.",
    )

    add_heading2(doc, "Email Alerts (Resend)")
    add_body(
        doc,
        "The email alerts module (email_alerts.py) sends notifications via the Resend API "
        "for critical errors, circuit breaker trips, generation failures, and daily digest "
        "summaries. All emails use clean HTML formatting with contextual styling (red for "
        "errors, neutral for digests). Rate limiting caps output at 10 emails per hour, "
        "with deduplication suppressing identical error_type+message combinations within "
        "a 30-minute window. The module is entirely disabled (no-op) when the Resend API "
        "key is not set.",
    )

    add_heading2(doc, "Health Endpoints & Dependency Probes")

    add_bullet(
        doc,
        " Deep health checks combining liveness, readiness, and dependency "
        "probes in a single response.",
        bold_prefix="Health API:",
    )
    add_bullet(
        doc,
        " Structured request metrics including latency, error rates, and "
        "throughput in a 1-hour rolling window.",
        bold_prefix="Metrics:",
    )
    add_bullet(
        doc,
        " Runtime tracking of memory consumption and disk usage for the "
        "data and cache directories.",
        bold_prefix="Resource Tracking:",
    )
    add_bullet(
        doc,
        " Reachability probes against external API dependencies to detect "
        "outages before they impact generation.",
        bold_prefix="Dependency Probes:",
    )

    add_heading2(doc, "SLO Monitoring & Error Budgets")
    add_body(
        doc,
        "The monitoring system tracks Service Level Objectives and error budget consumption. "
        "This enables proactive alerting when error rates approach thresholds, rather than "
        "reacting to incidents after they impact users. Audit trails record data "
        "transformation decisions for post-incident analysis.",
    )

    add_page_break(doc)

    # ===================================================================
    # PAGE 9: SECURITY & RELIABILITY
    # ===================================================================

    add_heading1(doc, "Security & Reliability")

    add_heading2(doc, "Input Sanitization")
    add_body(
        doc,
        "All user-facing inputs are sanitized before processing. Nova enforces a maximum "
        "message length of 4,000 characters and limits conversation history to 6 turns "
        "to prevent history injection attacks. Input validation occurs at every entry point "
        "(web UI, chat API, Slack webhook).",
    )

    add_heading2(doc, "Rate Limiting & Authentication")

    add_bullet(
        doc,
        " Configurable per-IP rate limiting prevents abuse of generation "
        "and chat endpoints.",
        bold_prefix="Per-IP Rate Limits:",
    )
    add_bullet(
        doc,
        " Global rate limiting on chat endpoints ensures fair usage across "
        "all concurrent users.",
        bold_prefix="Global Chat Limits:",
    )
    add_bullet(
        doc,
        " Administrative endpoints (QC results, metrics, health details) are "
        "protected by API key authentication.",
        bold_prefix="Admin API Key:",
    )
    add_bullet(
        doc,
        " Origin-based whitelisting restricts cross-origin requests to "
        "approved domains only.",
        bold_prefix="CORS Whitelisting:",
    )

    add_heading2(doc, "Adversarial Prompt Defense")
    add_body(
        doc,
        "The Nova chatbot includes adversarial prompt blocking that defends against five "
        "categories of prompt injection attacks. All five attack vectors tested during "
        "security evaluation were successfully blocked (5/5), preventing attempts to "
        "extract system prompts, bypass safety boundaries, or manipulate response behavior.",
    )

    add_heading2(doc, "Thread-Safe Architecture")
    add_body(
        doc,
        "All shared state in the application is protected by threading locks. This includes "
        "the LLM router's per-provider circuit breaker state, rate limit counters, the "
        "data orchestrator's lazy-loaded module references, Nova's conversation history, "
        "file I/O operations for caching and logging, and the Supabase cache layer. The "
        "application serves concurrent requests safely without data races or corruption.",
    )

    add_heading2(doc, "Circuit Breakers")
    add_body(
        doc,
        "Circuit breakers are deployed at two levels: per-provider circuit breakers in the "
        "LLM router (5 failures trigger 60-second cooldown per provider) and per-API circuit "
        "breakers in the enrichment layer (preventing cascading failures when external "
        "services are degraded). Circuit breaker state transitions trigger email alerts "
        "via the Resend integration.",
    )

    add_heading2(doc, "Graceful Shutdown")
    add_body(
        doc,
        "The monitoring module coordinates graceful shutdown across all background threads. "
        "When a shutdown signal is received, in-progress requests are allowed to complete, "
        "background health checks and QC runs are stopped cleanly, and any pending log "
        "shipments to Grafana Loki are flushed before the process exits.",
    )

    add_page_break(doc)

    # ===================================================================
    # PAGE 10: ARCHITECTURE SUMMARY
    # ===================================================================

    add_heading1(doc, "Architecture Summary")

    add_heading2(doc, "System Architecture")

    # Text-based architecture diagram
    arch_p = doc.add_paragraph()
    arch_p.paragraph_format.space_before = Pt(6)
    arch_p.paragraph_format.space_after = Pt(6)
    arch_text = (
        "                    +---------------------------+\n"
        "                    |        Web Browser         |\n"
        "                    |  (Excel/PPT + Nova Chat)   |\n"
        "                    +------------+--------------+\n"
        "                                 |\n"
        "              +------------------+------------------+\n"
        "              |                  |                  |\n"
        "     +--------v------+  +--------v------+  +-------v--------+\n"
        "     |   Web UI /    |  |  Nova Chat    |  |  Slack Bot     |\n"
        "     |  Generation   |  |  (nova.py)    |  | (nova_slack.py)|\n"
        "     |   (app.py)    |  |               |  |                |\n"
        "     +--------+------+  +--------+------+  +-------+--------+\n"
        "              |                  |                  |\n"
        "              +------------------+------------------+\n"
        "                                 |\n"
        "                    +------------v--------------+\n"
        "                    |   Data Orchestrator       |\n"
        "                    |  (data_orchestrator.py)   |\n"
        "                    +------------+--------------+\n"
        "                                 |\n"
        "         +-----------+-----------+-----------+-----------+\n"
        "         |           |           |           |           |\n"
        "   +-----v----+ +---v------+ +--v-------+ +-v--------+ +v---------+\n"
        "   | research | | trend_   | | collar_  | | 30+ Live | | data_    |\n"
        "   |   .py    | | engine   | | intel    | |   APIs   | |synthesize|\n"
        "   +----------+ +----------+ +----------+ +----------+ +----------+\n"
        "                                 |\n"
        "                    +------------v--------------+\n"
        "                    |     Budget Engine         |\n"
        "                    |   (budget_engine.py)      |\n"
        "                    +---------------------------+\n"
    )
    run = arch_p.add_run(arch_text)
    run.font.size = Pt(7.5)
    run.font.name = "Courier New"
    run.font.color.rgb = DARK_GRAY

    add_heading2(doc, "Technology Stack")

    stack_rows = [
        ["Language", "Python 3.11+"],
        [
            "Core Dependencies",
            "Zero external pip packages (stdlib only for server core)",
        ],
        ["HTTP Server", "http.server (Python stdlib)"],
        ["Concurrency", "threading + concurrent.futures.ThreadPoolExecutor"],
        ["LLM Integration", "urllib.request to 12 provider REST APIs"],
        ["Caching", "3-tier: in-memory dict, disk JSON, Supabase Postgres"],
        ["Logging", "Python logging + Grafana Cloud Loki (JSON structured)"],
        ["Alerting", "Resend email API + Slack webhooks"],
        ["Monitoring", "UptimeRobot (external) + custom health endpoints"],
        ["Deployment", "Render.com (Standard tier, persistent disk)"],
        ["Database", "Supabase (Postgres) for persistent L3 cache"],
    ]
    add_styled_table(
        doc,
        ["Component", "Technology"],
        stack_rows,
        col_widths=[Inches(1.5), Inches(4.5)],
    )

    add_heading2(doc, "Integration Ecosystem")

    integration_rows = [
        [
            "LLM Providers",
            "Gemini, Groq, Cerebras, Mistral, OpenRouter, xAI, SambaNova, NVIDIA NIM, Cloudflare, OpenAI, Anthropic",
        ],
        [
            "Government Data",
            "BLS (OES, QCEW, JOLTS), Census ACS, FRED, O*NET, CareerOneStop, H-1B",
        ],
        [
            "International Data",
            "World Bank, IMF, Eurostat, ILO, REST Countries, GeoNames, Teleport",
        ],
        [
            "Ad Platforms",
            "Google Ads, Meta Marketing, LinkedIn Marketing, Bing Ads, TikTok Marketing",
        ],
        ["Job Market", "Adzuna, Jooble, Indeed (via trends), DataUSA"],
        ["Company Intel", "Clearbit, Wikipedia, SEC EDGAR, Google Trends"],
        ["Observability", "Grafana Cloud Loki, UptimeRobot, Resend (email)"],
        ["Collaboration", "Slack (Bot + webhooks), Supabase (persistent cache)"],
    ]
    add_styled_table(
        doc,
        ["Category", "Integrations"],
        integration_rows,
        col_widths=[Inches(1.5), Inches(4.5)],
    )

    add_heading2(doc, "Future Roadmap")

    add_bullet(
        doc,
        " Expand Nova to support multi-turn media plan creation directly "
        "through conversational interaction.",
        bold_prefix="Conversational Plan Builder:",
    )
    add_bullet(
        doc,
        " Build a real-time dashboard showing plan performance, budget "
        "pacing, and market condition changes.",
        bold_prefix="Performance Analytics Dashboard:",
    )
    add_bullet(
        doc,
        " Integrate real-time bidding data to feed actual campaign "
        "performance back into future plan recommendations.",
        bold_prefix="Closed-Loop Optimization:",
    )
    add_bullet(
        doc,
        " Extend the evaluation framework with automated A/B testing of "
        "budget allocation strategies across historical campaigns.",
        bold_prefix="Automated A/B Testing Framework:",
    )
    add_bullet(
        doc,
        " Add Microsoft Teams integration alongside the existing "
        "Slack bot for broader workplace coverage.",
        bold_prefix="Teams Bot Integration:",
    )

    # --- Footer note ---
    doc.add_paragraph()
    sep_p2 = doc.add_paragraph()
    sep_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep_p2.add_run("_" * 60)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(10)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Confidential -- Shubham Singh Chandel | March 2026")
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    run.font.name = "Calibri"
    run.italic = True

    # --- Save ---
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(f"Document saved to: {path}")
    # Verify file size
    size_kb = os.path.getsize(path) / 1024
    print(f"File size: {size_kb:.1f} KB")
